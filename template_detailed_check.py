import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

# 打开模板文件
doc = docx.Document('template.docx')

print("=== Template.docx 详细检查报告 ===")
print()

# 1. 提取所有Jinja2标签
print("1. 所有Jinja2标签提取：")
all_tags = set()

for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text or '{%' in text:
        tags = re.findall(r'(\{\{[^}]+\}\}|\{\%[^%]+\%\})', text)
        if tags:
            print(f"   段落 {i+1}: {text.strip()}")
            for tag in tags:
                all_tags.add(tag.strip())

print(f"\n   总共找到 {len(all_tags)} 种不同的标签：")
for tag in sorted(all_tags):
    print(f"      - {tag}")

# 2. 表格分析
print("\n2. 表格分析：")
for i, table in enumerate(doc.tables):
    print(f"\n   表格 {i+1}：{len(table.rows)} 行 × {len(table.columns)} 列")
    
    # 查看表头
    if table.rows:
        header_cells = []
        for cell in table.rows[0].cells:
            header_cells.append(cell.text.strip())
        print(f"      表头：{header_cells}")
    
    # 查看前3行数据行（如果有）
    for row_idx in range(1, min(4, len(table.rows))):
        row_cells = []
        for cell in table.rows[row_idx].cells:
            row_cells.append(cell.text.strip())
        print(f"      行 {row_idx+1}：{row_cells}")
    
    # 检查是否有循环标签
    has_loop_tags = False
    for row in table.rows:
        for cell in row.cells:
            if '{% tr for' in cell.text or '{% endtr' in cell.text:
                has_loop_tags = True
                print(f"      包含循环标签：{cell.text.strip()}")
    if not has_loop_tags:
        print(f"      未发现循环标签")

# 3. 检查缺失的核心标签
print("\n3. 核心标签检查：")
required_static_tags = {'{{ company_name }}', '{{ report_year }}', '{{ total_emission }}'}
required_dynamic_tags = {'{% tr for item in items %}', '{% endtr %}', '{{ item.name }}', '{{ item.emission }}', '{{ item.note }}'}

missing_static = required_static_tags - all_tags
missing_dynamic = required_dynamic_tags - all_tags

if missing_static:
    print(f"   ❌ 缺少静态标签：{missing_static}")
else:
    print("   ✅ 所有静态标签已找到")

if missing_dynamic:
    print(f"   ❌ 缺少动态标签：{missing_dynamic}")
else:
    print("   ✅ 所有动态标签已找到")

# 4. 建议的修改位置
print("\n4. 建议的修改位置：")
print("   - 公司名称：在报告封面或首页找到公司名称位置，替换为 {{ company_name }}")
print("   - 报告年份：在报告标题或页眉页脚找到年份位置，替换为 {{ report_year }}")
print("   - 总排放量：在报告摘要或排放总量章节找到具体数值，替换为 {{ total_emission }}")
print("   - 动态表格：找到'减排行动列表'或'主要排放源列表'表格，按照要求添加循环标签")

print("\n=== 检查完成 ===")