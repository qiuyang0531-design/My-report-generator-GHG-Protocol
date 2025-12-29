#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
全面检查template.docx的问题
"""

import docx
import re

def comprehensive_check():
    """全面检查模板问题"""
    doc = docx.Document('template.docx')

    print("=== template.docx 全面问题检查 ===\n")

    # 1. 收集所有包含Jinja2标签的内容
    all_jinja_content = []
    variables_with_issues = []

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if '{{' in text or '{%' in text:
            all_jinja_content.append((f"段落{i+1}", text))
            # 查找所有变量
            vars_found = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
            for var in vars_found:
                var_name = var.strip()
                if any(issue in var_name for issue in [' ', '-']):
                    variables_with_issues.append((f"段落{i+1}", var_name))

    # 检查表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if '{{' in text or '{%' in text:
                        location = f"表格{table_idx+1} 行{row_idx+1} 列{col_idx+1}"
                        all_jinja_content.append((location, text))
                        vars_found = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                        for var in vars_found:
                            var_name = var.strip()
                            if any(issue in var_name for issue in [' ', '-']):
                                variables_with_issues.append((location, var_name))

    print("1. 发现的所有Jinja2变量问题:")
    if variables_with_issues:
        for location, var in variables_with_issues:
            print(f"   {location}: {var}")
    else:
        print("   未发现变量问题")

    print(f"\n2. 需要修复的具体问题总结:")

    # 提取唯一的问题变量
    unique_vars = set()
    for _, var in variables_with_issues:
        unique_vars.add(var)

    fixes = []
    for var in sorted(unique_vars):
        if ' ' in var:
            # 去掉空格
            fixed = re.sub(r'\s+', '_', var.strip())
            fixes.append((var, fixed))
        elif '-' in var:
            # 替换连字符
            fixed = var.replace('-', '_')
            fixes.append((var, fixed))

    print(f"   总共需要修复 {len(fixes)} 个变量:")
    for original, fixed in fixes:
        print(f"   - {{{{{original}}}}} → {{{{{fixed}}}}}")

    print(f"\n3. 当前数据准备代码中缺失的变量:")
    print("   根据之前的错误，以下变量可能需要添加:")
    print("   - scope_2_location (对应模板中的scope_2_location-based_emissions)")
    print("   - scope_2_market (对应模板中的scope_2_market-based_emissions)")
    print("   - scope_of_business (从{{ scope_of_ business }}修复而来)")

    print(f"\n4. 修复建议:")
    print("   a) 修复模板中的变量名（去除空格和连字符）")
    print("   b) 在数据准备代码中添加缺失的变量")
    print("   c) 或者创建一个兼容性层，将现有数据映射到模板变量")

    return fixes, variables_with_issues

if __name__ == "__main__":
    fixes, issues = comprehensive_check()