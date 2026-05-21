# main.py
# 使用重构后的协议驱动型数据读取器（方式1: 从新包导入）
from data_reader import ExcelDataReaderRefactored as ExcelDataReader
from docxtpl import DocxTemplate
from jinja2 import Environment
from datetime import date
import os
import re
from docx.oxml import OxmlElement
from inventory_summary_generator import generate_inventory_context
from report_config import get_scope_3_category_name, get_scope_3_category_name_by_chinese_numeral
from post_processor import (
    find_table_by_content,
    apply_chemical_subscripts,
    _insert_category1_material_subtitle,
    convert_parentheses_to_chinese,
    clean_excessive_blank_lines,
    insert_toc_field,
    fix_scope3_category_headers,
    add_excluded_categories_statement,
    clean_empty_category_tables_v2,
    merge_vertical_cells,
    _center_table_cells_horizontal,
    merge_other_tables_vertical_cells,
)

DEFAULT_DY_XLSX_NAME = "blank_form.xlsx"


def _extract_version_date(filename):
    """从文件名中提取版本日期，如 Update 20260518Protocol → 20260518"""
    import re
    m = re.search(r'Update\s*(\d{8})Protocol', filename)
    if m:
        return m.group(1)
    return "00000000"


def _find_latest_inventory_xlsx(search_dirs):
    """在搜索目录中按文件名版本日期查找最新的盘查清册 xlsx 文件"""
    candidates = []
    for d in search_dirs:
        if not d or not os.path.isdir(d):
            continue
        try:
            for name in os.listdir(d):
                if not name.lower().endswith(".xlsx"):
                    continue
                if name.startswith("~$"):
                    continue
                if "盘查清册" not in name:
                    continue
                if not name.startswith("DY-GHG-"):
                    continue
                full_path = os.path.join(d, name)
                candidates.append((_extract_version_date(name), full_path))
        except OSError:
            continue

    candidates.sort(key=lambda x: x[0], reverse=True)
    return candidates[0][1] if candidates else None


def resolve_inventory_xlsx_path(xlsx_path: str) -> str:
    home_dir = os.path.expanduser("~")
    search_dirs = [
        os.getcwd(),
        os.path.join(home_dir, "Desktop", "2026组织碳收资"),
        os.path.join(home_dir, "Desktop", "2026组织碳收资", "更新文件"),
    ]

    is_default = xlsx_path == DEFAULT_DY_XLSX_NAME
    latest = _find_latest_inventory_xlsx(search_dirs) if is_default else None

    if xlsx_path and os.path.exists(xlsx_path):
        if not is_default or not latest:
            return xlsx_path
        xlsx_date = _extract_version_date(os.path.basename(xlsx_path))
        latest_date = _extract_version_date(os.path.basename(latest))
        if xlsx_date >= latest_date:
            return xlsx_path

    if is_default and latest:
        return latest

    raise FileNotFoundError(
        f"Excel 文件不存在: {xlsx_path}"
        + (f"（已搜索目录: {search_dirs}）" if is_default else "")
    )



def to_chinese_num(n):
    """

    Args:
        n: 数字 (1-15)

    Returns:
        中文大写数字字符串
    """
    chinese_map = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五'
    }
    return chinese_map.get(n, str(n))


def format_number(value, decimals=2, with_comma=True):
    """
    格式化数字：添加千分位分隔符，保留指定小数位数（展示层格式化）

    Args:
        value: 数值（float 或 int）
        decimals: 小数位数，默认 2
        with_comma: 是否添加千分位分隔符，默认 True

    Returns:
        格式化后的字符串
    """
    try:
        float_value = float(value)
        if with_comma:
            return f"{float_value:,.{decimals}f}"
        else:
            return f"{float_value:.{decimals}f}"
    except (ValueError, TypeError):
        return "0.00"


def _init_protocol_vars(formatted_context):
    """初始化所有协议变量，确保 Jinja2 模板中可用"""
    protocol_output_vars = {
        'pro_ef_items': [],
        'emission_factor_items': [],
        'scope1_stationary_combustion_emissions_items': [],
        'scope1_mobile_combustion_emissions_items': [],
        'scope1_fugitive_emissions_items': [],
        'scope1_process_emissions_items': [],
        'gwp_items': [],
        'ghg_inventory_items': [],
        'activity_summary_items': [],
        'uncertainty_items': [],
        'reduction_action_items': [],
    }
    for var_name, default_value in protocol_output_vars.items():
        if var_name not in formatted_context:
            formatted_context[var_name] = default_value


def _clean_context_strings(d):
    """递归清洗字典中的所有字符串值：去首尾空格，合并多余空白，保留 \\n\\n 段落分隔"""
    if not isinstance(d, dict):
        return d
    cleaned = {}
    for key, value in d.items():
        if isinstance(value, str):
            v = str(value).strip()
            v = v.replace('\r\n', '\n').replace('\r', '\n')
            v = v.replace('\n\n', '\x00')
            v = re.sub(r'\s+', ' ', v)
            v = v.replace('\x00', '\n\n')
            cleaned[key] = v
        elif isinstance(value, dict):
            cleaned[key] = _clean_context_strings(value)
        elif isinstance(value, list):
            cleaned[key] = [
                _clean_context_strings(item) if isinstance(item, dict) else
                (re.sub(r'\s+', ' ', str(item).strip()) if isinstance(item, str) else item)
                for item in value
            ]
        else:
            cleaned[key] = value
    return cleaned

def _compute_scope3_sums(context, formatted_context):
    """计算范围三各类别排放总和及详细项列总和"""
    scope3_total_sum = 0
    for i in range(1, 16):
        cat_emission = context.get(f'scope_3_category_{i}_emissions', 0)
        if cat_emission and cat_emission > 0:
            scope3_total_sum += cat_emission

    formatted_context['scope3_categories_total_sum'] = scope3_total_sum
    formatted_context['scope3_categories_total_sum_formatted'] = format_number(scope3_total_sum)

    emission_columns = [
        'total_green_house_gas_emissions',
        'CO2_emissions', 'CH4_emissions', 'N2O_emissions',
        'HFCs_emissions', 'PFCs_emissions', 'SFs_emissions', 'NF3_emissions'
    ]

    column_key_mapping = {'SFs_emissions': 'SF6_emissions'}

    for cat_num in range(1, 16):
        category_var = f'scope3_category{cat_num}'
        if category_var in context and context[category_var]:
            column_sums = {col: 0.0 for col in emission_columns}
            for item in context[category_var]:
                for col in emission_columns:
                    data_key = column_key_mapping.get(col, col)
                    val = item.get(data_key, 0)
                    if val is None:
                        val = 0
                    try:
                        if isinstance(val, str):
                            val = float(val.replace(',', '').replace(' ', ''))
                        else:
                            val = float(val)
                        column_sums[col] += val
                    except (ValueError, TypeError):
                        pass

            for col in emission_columns:
                sum_value = column_sums[col]
                formatted_context[f'scope3_category{cat_num}_{col}_sum'] = sum_value
                formatted_context[f'scope3_category{cat_num}_{col}_sum_formatted'] = format_number(sum_value)


def _format_top_level_numbers(context, formatted_context):
    """格式化顶层数值字段、别名映射、范围三display字段及未纳入类别汇总"""
    number_fields = [
        'scope_1_emissions', 'scope_2_location_based_emissions',
        'scope_2_market_based_emissions', 'scope_3_emissions',
        'total_emission_location', 'total_emission_market',
    ]
    for i in range(1, 16):
        number_fields.append(f'scope_3_category_{i}_emissions')

    alias_fields = ['scope_1', 'scope_2_location', 'scope_2_market', 'scope_3']
    all_number_fields = number_fields + alias_fields

    field_aliases = {
        'scope_1': 'scope_1_emissions',
        'scope_2_location': 'scope_2_location_based_emissions',
        'scope_2_market': 'scope_2_market_based_emissions',
        'scope_3': 'scope_3_emissions',
    }

    for short_name, long_name in field_aliases.items():
        if short_name in context:
            formatted_context[long_name] = context[short_name]
            formatted_context[f'{long_name}_formatted'] = format_number(context[short_name])

    for field in all_number_fields:
        formatted_key = f'{field}_formatted'
        if formatted_key in formatted_context:
            continue
        if field in context and context[field] is not None:
            formatted_context[formatted_key] = format_number(context[field])
        else:
            formatted_context[formatted_key] = format_number(0)

    for i in range(1, 16):
        field = f'scope_3_category_{i}_emissions'
        value = context.get(field, 0)
        display_value = format_number(value) if value and value > 0 else ""
        formatted_context[f'{field}_display'] = display_value
        formatted_context[f'category_{i}_emissions_display'] = display_value

    categories_not_in_scope = []
    for i in range(1, 16):
        field = f'scope_3_category_{i}_emissions'
        value = context.get(field, 0)
        if not value or value <= 0:
            categories_not_in_scope.append(i)

    if categories_not_in_scope:
        category_str = "、".join(f"类别{cat}" for cat in categories_not_in_scope)
        formatted_context['scope_3_categories_not_in_scope_summary'] = (
            f"范围三{category_str}产生的排放为0，排放量不在本次盘查范围内，本周期内不进行量化。"
        )
    else:
        formatted_context['scope_3_categories_not_in_scope_summary'] = ""


def _format_category_numbers(categories):
    """将类别数字列表格式化为合并字符串，如 [1,2,3,4,5,6,7,9,10,11,12] -> "1~7、9~12" """
    if not categories:
        return ""
    sorted_cats = sorted(categories)
    result = []
    start = end = sorted_cats[0]
    for i in range(1, len(sorted_cats)):
        if sorted_cats[i] == end + 1:
            end = sorted_cats[i]
        else:
            result.append(str(start) if start == end else f"{start}~{end}")
            start = end = sorted_cats[i]
    result.append(str(start) if start == end else f"{start}~{end}")
    return "、".join(result)


def _generate_boundary_text(formatted_context, context):
    """动态生成盘查边界描述文本"""
    def _sf(v):
        try:
            return float(v) if v is not None else 0.0
        except (ValueError, TypeError):
            return 0.0

    parts = []
    if _sf(context.get('scope_1_emissions', 0)) > 0:
        parts.append("范围一")
    if _sf(context.get('scope_2_location_based_emissions', 0)) > 0 or \
       _sf(context.get('scope_2_market_based_emissions', 0)) > 0:
        parts.append("范围二")
    valid_cats = [i for i in range(1, 16)
                  if _sf(context.get(f'scope_3_category_{i}_emissions', 0)) > 0]
    if valid_cats:
        parts.append(f"范围三类别{_format_category_numbers(valid_cats)}")
    formatted_context['included_scopes_text'] = "、".join(parts) if parts else "所有范围"
    print(f"[盘查边界] 生成描述文本: {formatted_context['included_scopes_text']}")




def _format_activity_summary(context, formatted_context):
    """格式化活动数据汇总表（基于位置和基于市场）"""
    emission_fields = [
        'CO2_emissions', 'CH4_emissions', 'N2O_emissions',
        'HFCs_emissions', 'PFCs_emissions', 'SF6_emissions', 'NF3_emissions',
        'total_green_house_gas_emissions'
    ]

    def _fmt_one(variant_key, act_summary_key, activity_data_key, label):
        """格式化单个活动数据汇总表变体"""
        if act_summary_key in context and context[act_summary_key]:
            formatted_rows = []
            fields_to_fmt = [act_summary_key, activity_data_key] + emission_fields
            sums = {f: 0.0 for f in emission_fields}

            for item in context[act_summary_key]:
                fmt_item = item.copy()
                for field in fields_to_fmt:
                    if field in item:
                        orig = item[field]
                        try:
                            num = float(str(orig).replace(',', '').replace(' ', '')) if orig else 0
                            if field in sums:
                                sums[field] += num
                        except (ValueError, TypeError):
                            pass
                        if field == act_summary_key:
                            fmt_item[field] = format_number(orig) if orig else '0.00'
                        elif field == activity_data_key:
                            fmt_item[f'{field}_formatted'] = format_number(orig) if orig else '0.00'
                        else:
                            fmt_item[field] = format_number(orig) if orig else '0.00'
                formatted_rows.append(fmt_item)

            formatted_context[act_summary_key] = formatted_rows
            print(f"[活动数据汇总表] 已格式化 {len(formatted_rows)} 行数据{label}")

            for f in emission_fields:
                formatted_context[f'{variant_key}_{f}_sum_formatted'] = format_number(sums[f])
            print(f"[活动数据汇总表] 汇总行计算完成{label}")
        else:
            formatted_context[act_summary_key] = []
            for f in emission_fields:
                formatted_context[f'{variant_key}_{f}_sum_formatted'] = '0.00'

    _fmt_one('loc', 'act_summary_loc', 'activity_data_location_based', '')
    _fmt_one('mar', 'act_summary_mar', 'activity_data_market_based', '（基于市场）')


def _format_emission_factors(context, formatted_context):
    """格式化排放因子汇总表（pro_ef_items、indir_ef_items 及范围三各类别 EF）"""
    # --- pro_ef_items ---
    if 'pro_ef_items' in context and context['pro_ef_items']:
        formatted_items = []
        ef_fields = ['ncv', 'ox_rate', 'ef_val',
                     'CO2_emission_factor', 'CH4_emission_factor', 'N2O_emission_factor']
        ef_sums = {f: 0.0 for f in ef_fields}
        ef_count = 0

        for item in context['pro_ef_items']:
            fmt_item = item.copy()
            is_valid = False

            if 'number' in item:
                v = item['number']
                try:
                    v_str = str(v)
                    fmt_item['number'] = str(int(float(v))) if v_str.endswith('.0') or re.match(r'^\d+\.0$', v_str) else v_str
                except (ValueError, TypeError):
                    fmt_item['number'] = str(v) if v is not None else ""

            for field in ef_fields:
                if field in item:
                    orig = item[field]
                    try:
                        num = float(str(orig).replace(',', '').replace(' ', '')) if orig else 0
                        if field in ef_sums:
                            ef_sums[field] += num
                            is_valid = True
                    except (ValueError, TypeError):
                        pass
                    fmt_item[field] = format_number(orig) if orig else '0.00'

            if is_valid:
                ef_count += 1
            formatted_items.append(fmt_item)

        formatted_context['pro_ef_items'] = formatted_items
        formatted_context['emission_factor_items'] = formatted_items
        print(f"[排放因子汇总表] 已格式化 {len(formatted_items)} 行数据")

        if ef_count > 0:
            for f in ef_fields:
                formatted_context[f'ef_{f}_sum_formatted'] = format_number(ef_sums[f] / ef_count)
        else:
            for f in ef_fields:
                formatted_context[f'ef_{f}_sum_formatted'] = '0.00'
        print(f"[排放因子汇总表] 汇总行计算完成（基于{ef_count}行有效数据）")
    else:
        formatted_context['pro_ef_items'] = []
        formatted_context['emission_factor_items'] = []
        for f in ['ncv', 'ox_rate', 'ef_val', 'CO2_emission_factor', 'CH4_emission_factor', 'N2O_emission_factor']:
            formatted_context[f'ef_{f}_sum_formatted'] = '0.00'

    # --- indir_ef_items ---
    if 'indir_ef_items' in context and context['indir_ef_items']:
        formatted_items = []
        for item in context['indir_ef_items']:
            fmt_item = item.copy()
            if 'number' in item:
                v = item['number']
                try:
                    v_str = str(v)
                    fmt_item['number'] = str(int(float(v))) if v_str.endswith('.0') or re.match(r'^\d+\.0$', v_str) else v_str
                except (ValueError, TypeError):
                    fmt_item['number'] = str(v) if v is not None else ""
            if 'elec_emission_factor' in item:
                fmt_item['elec_emission_factor'] = format_number(item['elec_emission_factor'], decimals=4)
            formatted_items.append(fmt_item)
        formatted_context['indir_ef_items'] = formatted_items
        print(f"[范围二排放因子] 已格式化 indir_ef_items: {len(formatted_items)} 行数据")

    # --- 范围三各类别 EF (cat1-cat15) ---
    for i in range(1, 16):
        var_name = f'cat{i}_ef_items'
        if var_name in context and context[var_name]:
            formatted_items = []
            for item in context[var_name]:
                fmt_item = item.copy()
                for k, v in item.items():
                    if k == 'number':
                        try:
                            fmt_item[k] = str(int(float(v))) if v is not None else ""
                        except (ValueError, TypeError):
                            fmt_item[k] = str(v) if v is not None else ""
                    elif isinstance(v, (int, float)):
                        if k.endswith('_emission_factor'):
                            fmt_item[k] = format_number(v, decimals=5, with_comma=False)
                        else:
                            fmt_item[k] = format_number(v)
                formatted_items.append(fmt_item)
            formatted_context[var_name] = formatted_items
            print(f"[范围三排放因子] 已格式化 {var_name}: {len(formatted_items)} 行数据")



def _format_detail_tables(context, formatted_context):
    """格式化范围三、范围二、范围一排放明细表"""
    # --- 范围三排放明细表 (scope3_category1-15) ---
    for i in range(1, 16):
        var_name = f'scope3_category{i}'
        if var_name in context and context[var_name]:
            formatted_items = []
            for item in context[var_name]:
                formatted_item = item.copy()
                for k, v in item.items():
                    if k == 'number':
                        try:
                            v_str = str(v)
                            if v_str.endswith('.0') or re.match(r'^\d+\.0$', v_str):
                                formatted_item[k] = str(int(float(v)))
                            else:
                                formatted_item[k] = v_str
                        except (ValueError, TypeError):
                            formatted_item[k] = str(v) if v is not None else ""
                    elif isinstance(v, (int, float)):
                        formatted_item[k] = format_number(v)
                if 'SF6_emissions' in formatted_item:
                    formatted_item['SFs_emissions'] = formatted_item['SF6_emissions']
                formatted_items.append(formatted_item)
            formatted_context[var_name] = formatted_items
            print(f"[范围三明细] 已格式化 {var_name}: {len(formatted_items)} 行数据")

    # --- 范围二排放明细表 (scope2_items) ---
    if 'scope2_items' in context and context['scope2_items']:
        formatted_items = []
        for item in context['scope2_items']:
            formatted_item = item.copy()
            for k, v in item.items():
                if k == 'number':
                    try:
                        v_str = str(v)
                        if v_str.endswith('.0') or re.match(r'^\d+\.0$', v_str):
                            formatted_item[k] = str(int(float(v)))
                        else:
                            formatted_item[k] = v_str
                    except (ValueError, TypeError):
                        formatted_item[k] = str(v) if v is not None else ""
                elif isinstance(v, (int, float)):
                    formatted_item[k] = format_number(v)
            if 'SF6_emissions' in formatted_item:
                formatted_item['SFs_emissions'] = formatted_item['SF6_emissions']
            formatted_items.append(formatted_item)
        formatted_context['scope2_items'] = formatted_items
        print(f"[范围二明细] 已格式化 scope2_items: {len(formatted_items)} 行数据")

    # --- 范围一直接排放源清册数据 ---
    scope1_vars = [
        'scope1_stationary_combustion_emissions_items',
        'scope1_mobile_combustion_emissions_items',
        'scope1_fugitive_emissions_items',
        'scope1_process_emissions_items',
    ]
    emission_fields = [
        'CO2_emissions', 'CH4_emissions', 'N2O_emissions',
        'HFCs_emissions', 'PFCs_emissions', 'SF6_emissions', 'NF3_emissions',
        'total_green_house_gas_emissions'
    ]

    for var_name in scope1_vars:
        if var_name in context and context[var_name]:
            formatted_items = []
            for item in context[var_name]:
                formatted_item = item.copy()
                if 'number' in item:
                    v = item['number']
                    try:
                        v_str = str(v)
                        if v_str.endswith('.0') or re.match(r'^\d+\.0$', v_str):
                            formatted_item['number'] = str(int(float(v)))
                        else:
                            formatted_item['number'] = v_str
                    except (ValueError, TypeError):
                        formatted_item['number'] = str(v) if v is not None else ""
                for field in emission_fields:
                    if field in item:
                        original_value = item[field]
                        formatted_value = format_number(original_value) if original_value is not None else '0.00'
                        formatted_item[field] = formatted_value
                        if field == 'SF6_emissions':
                            formatted_item['SFs_emissions'] = formatted_value
                formatted_items.append(formatted_item)
            formatted_context[var_name] = formatted_items
            print(f"[范围一排放] 已格式化 {var_name}: {len(formatted_items)} 行数据")
        else:
            formatted_context[var_name] = []


def _final_string_clean(d):
    """递归清洗所有字符串值，去除首尾空格"""
    if isinstance(d, dict):
        return {k: _final_string_clean(v) for k, v in d.items()}
    elif isinstance(d, list):
        return [_final_string_clean(item) for item in d]
    elif isinstance(d, str):
        return d.strip()
    else:
        return d


def _finalize_context(formatted_context, context):
    """最终处理：cn_nums 映射、盘查边界、字符串清洗、总排放计算"""
    formatted_context['cn_nums'] = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
        16: '十六', 17: '十七', 18: '十八', 19: '十九', 20: '二十',
        21: '二十一', 22: '二十二', 23: '二十三', 24: '二十四', 25: '二十五',
        26: '二十六', 27: '二十七', 28: '二十八', 29: '二十九', 30: '三十'
    }

    _generate_boundary_text(formatted_context, context)

    # 最终字符串清洗
    formatted_context = _final_string_clean(formatted_context)

    # 总温室气体排放量计算
    def _sf(v):
        try:
            return float(v) if v is not None else 0.0
        except (ValueError, TypeError):
            return 0.0

    location_based_total = (
        _sf(context.get('scope_1_emissions', 0)) +
        _sf(context.get('scope_2_location_based_emissions', 0)) +
        _sf(context.get('scope_3_emissions', 0))
    )
    fmt_loc = format_number(location_based_total)
    formatted_context['location_based_total_green_house_gas_emissions'] = fmt_loc
    formatted_context['location_based_total_green_house_gas_emissions_formatted'] = fmt_loc

    market_based_total = (
        _sf(context.get('scope_1_emissions', 0)) +
        _sf(context.get('scope_2_market_based_emissions', 0)) +
        _sf(context.get('scope_3_emissions', 0))
    )
    fmt_mkt = format_number(market_based_total)
    formatted_context['market_based_total_green_house_gas_emissions'] = fmt_mkt
    formatted_context['market_based_total_green_house_gas_emissions_formatted'] = fmt_mkt

    import sys
    print(f"[总排放量计算]", file=sys.stderr)
    print(f"  基于位置的总排放量: {format_number(location_based_total)} 吨CO2e", file=sys.stderr)
    print(f"  基于市场的总排放量: {format_number(market_based_total)} 吨CO2e", file=sys.stderr)

    return formatted_context

def prepare_context_with_formatting(context):
    """为 context 添加格式化后的数值版本，返回新字典"""
    formatted_context = context.copy()

    _init_protocol_vars(formatted_context)
    formatted_context = _clean_context_strings(formatted_context)

    _compute_scope3_sums(context, formatted_context)
    _format_top_level_numbers(context, formatted_context)
    _format_activity_summary(context, formatted_context)
    _format_emission_factors(context, formatted_context)

    _format_detail_tables(context, formatted_context)
    formatted_context = _finalize_context(formatted_context, context)
    return formatted_context

def _build_summary_data_base(workbook):
    """从Excel的"基准年温室气体清单"sheet读取基准年汇总数据。

    读取左半区（cols 0-9）的实际数值，构建与 summary_data 结构一致的字典。
    GHG位置 sheet → scope2.loc + total_loc
    GHG市场 sheet → scope2.mar + total_mar
    scope1 和 scope3 由两张表共享（数值相同）。
    """
    import openpyxl
    from inventory_summary_generator import format_number as fmt

    gases = ['co2', 'ch4', 'n2o', 'hfcs', 'pfcs', 'sf6', 'nf3', 'total']

    def _safe_float(v):
        if v is None:
            return 0.0
        if isinstance(v, str):
            v = v.replace(',', '')
        try:
            return float(v)
        except (ValueError, TypeError):
            return 0.0

    def _read_gas_row(row, start=2):
        return {g: _safe_float(row[start + i]) if start + i < len(row) else 0.0
                for i, g in enumerate(gases)}

    def _fmt_dict(d):
        return {k: fmt(v) for k, v in d.items()}

    def _find_base_year_sheets(wb):
        """通过结构特征识别基准年清单sheet，返回 (ws_loc, ws_mar)。

        活动数据汇总 sheet 总是 37 列（固定格式），
        基准年清单 sheet 列数更少（11-23 列，视是否有右半区 Jinja2 标签而定）。
        两个基准年 sheet 中列数较多的为基于位置，较少的为基于市场。
        """
        candidates = []
        for ws in wb.worksheets:
            if 'GHG' not in ws.title:
                continue
            if ws.max_column == 37:  # 活动数据汇总，跳过
                continue
            candidates.append(ws)

        if len(candidates) < 2:
            return None, None

        # 按列数降序：多的为位置，少的为市场
        candidates.sort(key=lambda ws: ws.max_column, reverse=True)
        return candidates[0], candidates[1]

    def _is_data_row(row_vals):
        """判断一行是否为数据行：col[1] 非空，col[2]（CO2）为数值。"""
        if len(row_vals) < 10:
            return False
        col1 = str(row_vals[1]).strip() if row_vals[1] is not None else ''
        if not col1:
            return False
        # 排除表头行（col[2] = 'CO2' 等非数值文本）
        try:
            float(str(row_vals[2]).replace(',', ''))
            return True
        except (ValueError, TypeError):
            return False

    def _parse_sheet(ws):
        """内容感知解析：自动识别数据行，兼容不同行布局。

        返回 (scope1, scope2, cats_dict, scope3_total, grand_total)。
        """
        # 收集主表区所有数据行（row < 53），按出现顺序排列
        main_data = []
        for row in ws.iter_rows(min_row=1, max_row=52, values_only=True):
            if _is_data_row(row):
                main_data.append(_read_gas_row(row))

        scope1 = main_data[0] if len(main_data) > 0 else {}
        scope2 = main_data[1] if len(main_data) > 1 else {}
        cats = {}
        for i in range(15):
            idx = 2 + i
            if idx < len(main_data):
                cats[f'cat{i + 1}'] = main_data[idx]

        # 合计区（row 53+）：收集所有数据行
        summary_data = []
        for row in ws.iter_rows(min_row=53, values_only=True):
            if _is_data_row(row):
                summary_data.append(_read_gas_row(row))

        # 合计区按顺序：范围一、范围二、范围三、总计
        scope3_total = summary_data[2] if len(summary_data) > 2 else {}
        grand_total = summary_data[3] if len(summary_data) > 3 else {}

        return scope1, scope2, cats, scope3_total, grand_total

    ws_loc, ws_mar = _find_base_year_sheets(workbook)

    if not ws_loc or not ws_mar:
        print("[基准年清单] 未找到基准年温室气体清单sheet，跳过 summary_data_base")
        return None

    loc = _parse_sheet(ws_loc)
    mar = _parse_sheet(ws_mar)

    scope1, scope2_loc, cats_loc, scope3_loc_total, total_loc = loc
    _, scope2_mar, cats_mar, scope3_mar_total, total_mar = mar

    # scope1 和 scope3 以位置sheet为准（与市场sheet数值相同）
    scope3 = _fmt_dict(scope3_loc_total) if scope3_loc_total else {}
    for cat_key, cat_vals in cats_loc.items():
        scope3[cat_key] = _fmt_dict(cat_vals)

    result = {
        'scope1': _fmt_dict(scope1),
        'scope2': {
            'loc': _fmt_dict(scope2_loc),
            'mar': _fmt_dict(scope2_mar),
        },
        'scope3': scope3,
        'total': _fmt_dict(total_loc) if total_loc else {},
        'total_loc': _fmt_dict(total_loc) if total_loc else {},
        'total_mar': _fmt_dict(total_mar) if total_mar else {},
        'total_mkt': _fmt_dict(total_mar) if total_mar else {},
    }
    result['scope2']['mkt'] = result['scope2']['mar']

    print(f"[基准年清单] 基于位置总计: {result['total_loc']['total']} 吨CO2e")
    print(f"[基准年清单] 基于市场总计: {result['total_mar']['total']} 吨CO2e")
    return result


def generate_report_from_xlsx(
    xlsx_path=DEFAULT_DY_XLSX_NAME,
    template_path="blank_template.docx",
    output_path="carbon_report.docx"
):
    """
    使用 template.docx 作为模板，从 xlsx 文件动态读取数据生成报告
    数字格式化在展示层通过 Jinja2 过滤器处理，数据层保持原始数字类型

    Args:
        xlsx_path: Excel 数据文件路径
        template_path: Word 模板文件路径
        output_path: 输出报告路径（默认: carbon_report.docx）
    """
    print("=" * 50)
    print("开始生成碳盘查报告（纯xlsx，动态读取）")
    print("=" * 50)

    xlsx_path = resolve_inventory_xlsx_path(xlsx_path)

    # 1. 使用 data_reader 的协议驱动方法（重构版）
    print(f"\n[步骤1] 从 {xlsx_path} 动态提取数据...")
    reader = ExcelDataReader(xlsx_path)
    context = reader.get_all_context()  # 重构后使用 get_all_context()
    reader.close()  # 重构后需要手动关闭工作簿

    # 打印提取的关键数据
    print("\n提取的关键数据:")
    print(f"  公司名称: {context.get('company_name')}")
    print(f"  范围一排放: {context.get('scope_1_emissions')}")
    print(f"  范围二排放（基于位置）: {context.get('scope_2_location_based_emissions')}")
    print(f"  范围二排放（基于市场）: {context.get('scope_2_market_based_emissions')}")
    print(f"  范围三排放: {context.get('scope_3_emissions')}")

    # 2. 加载模板
    print(f"\n[步骤2] 加载模板: {template_path}")
    template = DocxTemplate(template_path)

    # 3. 准备展示层数据（格式化数字，保持数据层纯净）
    print("\n[步骤3] 准备展示层数据（格式化数字）...")
    render_context = prepare_context_with_formatting(context)

    # 3.5 生成基准年温室气体清单汇总数据
    print("\n[步骤3.5] 生成基准年温室气体清单汇总数据...")

    # ============================================================
    # 辅助函数：安全转换 + 过滤合计行
    # ============================================================
    def safe_float(value):
        """安全转换为浮点数"""
        try:
            # 移除千分位逗号
            if isinstance(value, str):
                value = value.replace(',', '')
            return float(value) if value else 0.0
        except (ValueError, TypeError):
            return 0.0

    def is_summary_row(item):
        """检查是否为合计/小计行，需要过滤掉"""
        # 检查多个可能的字段
        for field in ['emission_source', 'category', 'name', 'number']:
            value = str(item.get(field, '')).strip()
            # 检查是否包含合计、小计、总计等关键字
            if any(keyword in value for keyword in ['合计', '小计', '总计', '汇总', 'Total', 'Sum']):
                return True
        return False

    # ============================================================
    # 第一步：计算范围一的各项气体总和（从4个明细列表累加）
    # ============================================================
    scope1_lists = [
        'scope1_stationary_combustion_emissions_items',  # 固定燃烧
        'scope1_mobile_combustion_emissions_items',      # 移动燃烧
        'scope1_fugitive_emissions_items',               # 逸散排放
        'scope1_process_emissions_items',                # 制程排放
    ]

    # 初始化范围一累加器
    scope1_sums = {
        'co2': 0.0, 'ch4': 0.0, 'n2o': 0.0,
        'hfcs': 0.0, 'pfcs': 0.0, 'sf6': 0.0, 'nf3': 0.0, 'total': 0.0
    }

    # 遍历4个明细列表，累加各项气体（过滤掉合计行）
    filtered_count = 0
    for list_name in scope1_lists:
        items = context.get(list_name, [])
        for item in items:
            # 过滤掉合计/小计行
            if is_summary_row(item):
                filtered_count += 1
                continue
            scope1_sums['co2'] += safe_float(item.get('CO2_emissions'))
            scope1_sums['ch4'] += safe_float(item.get('CH4_emissions'))
            scope1_sums['n2o'] += safe_float(item.get('N2O_emissions'))
            scope1_sums['hfcs'] += safe_float(item.get('HFCs_emissions'))
            scope1_sums['pfcs'] += safe_float(item.get('PFCs_emissions'))
            scope1_sums['sf6'] += safe_float(item.get('SFs_emissions'))  # 注意是SFs不是SF6
            scope1_sums['nf3'] += safe_float(item.get('NF3_emissions'))
            scope1_sums['total'] += safe_float(item.get('total_green_house_gas_emissions'))

    # 使用context中的汇总值作为总计（更准确）
    scope1_total_from_context = safe_float(context.get('scope_1_emissions', 0))
    if scope1_total_from_context > 0:
        scope1_sums['total'] = scope1_total_from_context

    print(f"[范围一汇总] CO2: {scope1_sums['co2']:.2f}, CH4: {scope1_sums['ch4']:.2f}, 总计: {scope1_sums['total']:.2f}")
    if filtered_count > 0:
        print(f"[范围一汇总] 已过滤 {filtered_count} 个合计/小计行")

    # ============================================================
    # 第二步：构建 summary_raw_data
    # ============================================================
    # 提取范围二的两组数据（基于位置和基于市场）
    scope2_loc_co2 = safe_float(context.get('scope_2_location_based_emissions', 0))
    scope2_mar_co2 = safe_float(context.get('scope_2_market_based_emissions', 0))

    summary_raw_data = {
        # 范围一排放数据（使用计算出的总和）
        'scope1_co2': scope1_sums['co2'],
        'scope1_ch4': scope1_sums['ch4'],
        'scope1_n2o': scope1_sums['n2o'],
        'scope1_hfcs': scope1_sums['hfcs'],
        'scope1_pfcs': scope1_sums['pfcs'],
        'scope1_sf6': scope1_sums['sf6'],
        'scope1_nf3': scope1_sums['nf3'],
        'scope1_total': scope1_sums['total'],

        # 范围二排放数据 - 基于位置
        'scope2_loc_co2': scope2_loc_co2,
        'scope2_loc_ch4': 0.0,
        'scope2_loc_n2o': 0.0,
        'scope2_loc_hfcs': 0.0,
        'scope2_loc_pfcs': 0.0,
        'scope2_loc_sf6': 0.0,
        'scope2_loc_nf3': 0.0,
        'scope2_loc_total': scope2_loc_co2,

        # 范围二排放数据 - 基于市场
        'scope2_mar_co2': scope2_mar_co2,
        'scope2_mar_ch4': 0.0,
        'scope2_mar_n2o': 0.0,
        'scope2_mar_hfcs': 0.0,
        'scope2_mar_pfcs': 0.0,
        'scope2_mar_sf6': 0.0,
        'scope2_mar_nf3': 0.0,
        'scope2_mar_total': scope2_mar_co2,
    }

    # ============================================================
    # 第三步：添加范围三类别数据
    # 优先使用context中的汇总值，分项气体从明细累加（过滤合计行）
    # ============================================================
    for i in range(1, 16):
        cat_key = f'cat{i}'
        cat_list = context.get(f'scope3_category{i}', [])

        # 累加该类别的各项气体（过滤掉合计行）
        cat_sums = {
            'co2': 0.0, 'ch4': 0.0, 'n2o': 0.0,
            'hfcs': 0.0, 'pfcs': 0.0, 'sf6': 0.0, 'nf3': 0.0, 'total': 0.0
        }

        cat_filtered_count = 0
        for item in cat_list:
            # 过滤掉合计/小计行
            if is_summary_row(item):
                cat_filtered_count += 1
                continue
            cat_sums['co2'] += safe_float(item.get('CO2_emissions'))
            cat_sums['ch4'] += safe_float(item.get('CH4_emissions'))
            cat_sums['n2o'] += safe_float(item.get('N2O_emissions'))
            cat_sums['hfcs'] += safe_float(item.get('HFCs_emissions'))
            cat_sums['pfcs'] += safe_float(item.get('PFCs_emissions'))
            cat_sums['sf6'] += safe_float(item.get('SFs_emissions'))  # 注意是SFs
            cat_sums['nf3'] += safe_float(item.get('NF3_emissions'))
            cat_sums['total'] += safe_float(item.get('total_green_house_gas_emissions'))

        # 使用context中的汇总值作为总计（更准确）
        cat_total_from_context = safe_float(context.get(f'scope_3_category_{i}_emissions', 0))
        gas_sum = sum(cat_sums[g] for g in ['co2', 'ch4', 'n2o', 'hfcs', 'pfcs', 'sf6', 'nf3'])
        
        if cat_total_from_context > 0:
            cat_sums['total'] = cat_total_from_context
        elif gas_sum > 0:
            cat_sums['total'] = gas_sum
        
        # 确保总计不小于分项气体之和（修正可能的四舍五入误差）
        if gas_sum > cat_sums['total']:
            cat_sums['total'] = gas_sum

        summary_raw_data[cat_key] = cat_sums

    print(f"[范围三汇总] 类别1 CO2: {summary_raw_data['cat1']['co2']:.2f}, 总计: {summary_raw_data['cat1']['total']:.2f}")

    # 生成summary_data
    summary_data = generate_inventory_context(summary_raw_data)
    render_context['summary_data'] = summary_data
    print("[汇总数据] 已添加到渲染上下文")

    # 添加排放源明细数据到summary_data
    # ============================================================
    print("[步骤3.6] 添加排放源明细数据...")

    # 范围一排放源明细（固定燃烧、移动燃烧、逸散排放、制程排放）
    scope1_detail_lists = {
        'stationary': context.get('scope1_stationary_combustion_emissions_items', []),
        'mobile': context.get('scope1_mobile_combustion_emissions_items', []),
        'fugitive': context.get('scope1_fugitive_emissions_items', []),
        'process': context.get('scope1_process_emissions_items', []),
    }
    summary_data['scope1_detail'] = scope1_detail_lists
    print(f"  [范围一明细] 固定燃烧: {len(scope1_detail_lists['stationary'])} 条")
    print(f"  [范围一明细] 移动燃烧: {len(scope1_detail_lists['mobile'])} 条")
    print(f"  [范围一明细] 逸散排放: {len(scope1_detail_lists['fugitive'])} 条")
    print(f"  [范围一明细] 制程排放: {len(scope1_detail_lists['process'])} 条")

    # 范围二排放源明细（外购能源间接排放）
    summary_data['scope2_detail'] = {
        'loc': context.get('indir_ef_items', []),  # 基于位置
        'mkt': context.get('indir_ef_items', []),  # 基于市场（相同数据）
    }
    print(f"  [范围二明细] 外购能源: {len(summary_data['scope2_detail']['loc'])} 条")

    # 范围三排放源明细（15个类别）
    scope3_detail_lists = {}
    for i in range(1, 16):
        scope3_detail_lists[f'cat{i}'] = context.get(f'scope3_category{i}', [])
    summary_data['scope3_detail'] = scope3_detail_lists
    print(f"  [范围三明细] 共 {sum(len(v) for v in scope3_detail_lists.values())} 条明细")

    # 活动数据汇总明细
    summary_data['activity_summary_detail'] = {
        'loc': context.get('activity_summary_items', []),   # 基于位置
        'mkt': context.get('activity_summary_market_items', []),  # 基于市场
    }
    print(f"  [活动数据汇总] 基于位置: {len(summary_data['activity_summary_detail']['loc'])} 条")
    print(f"  [活动数据汇总] 基于市场: {len(summary_data['activity_summary_detail']['mkt'])} 条")

    # 更新render_context中的summary_data
    render_context['summary_data'] = summary_data

    # 3.5.5 从Excel读取基准年温室气体清单sheet，构建 summary_data_base
    print("\n[步骤3.5.5] 读取基准年温室气体清单sheet...")
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    try:
        summary_data_base = _build_summary_data_base(wb)
        if summary_data_base is not None:
            render_context['summary_data_base'] = summary_data_base
            print("[基准年清单] summary_data_base 已添加到渲染上下文")
        else:
            print("[基准年清单] 未生成 summary_data_base，模板中的 {{summary_data_base.*}} 将显示为空")
    finally:
        wb.close()

    # 3.6 拆分 scope_2 数据：区分基于位置和基于市场的 AD 描述
    quant_methods = render_context.get('quantification_methods', {})
    scope_2_all = quant_methods.get('scope_2', {})
    if scope_2_all:
        # scope_2 保留基于位置的条目（不含 market 后缀的 key）
        quant_methods['scope_2'] = {
            k: v for k, v in scope_2_all.items() if 'market' not in k.lower()
        }
        # scope_2_market 保留基于市场的条目（含 market 后缀的 key）
        quant_methods['scope_2_market'] = {
            k: v for k, v in scope_2_all.items() if 'market' in k.lower()
        }
        # 确保基于市场的 AD 提及绿色电力证书（GEC）
        for k, v in quant_methods['scope_2_market'].items():
            ad = v.get('ad', '')
            if '绿色电力证书' not in ad and 'GEC' not in ad:
                v['ad'] = f"在基于位置的基础上，扣除已核销的绿色电力证书（GEC）对应电量。{ad}"
        print("[范围二] 已拆分 scope_2 (基于位置) 和 scope_2_market (基于市场)")

    # 3.7 过滤范围三量化方法：仅保留有排放的类别
    flags = render_context.get('flags', {})
    scope_3_methods = quant_methods.get('scope_3', {})
    scope_3_names = render_context.get('scope_3_category_names', {})
    if scope_3_methods:
        active_scope_3 = {}
        active_names = {}
        for i in range(1, 16):
            cat_key = f'category_{i}'
            flag_key = f'has_scope_3_category_{i}'
            if flags.get(flag_key, False):
                if cat_key in scope_3_methods:
                    active_scope_3[cat_key] = scope_3_methods[cat_key]
                # scope_3_category_names 使用字符串键 "category_i"
                if cat_key in scope_3_names:
                    active_names[cat_key] = scope_3_names[cat_key]
                elif i in scope_3_names:
                    active_names[cat_key] = scope_3_names[i]
        quant_methods['scope_3'] = active_scope_3
        render_context['scope_3_category_names'] = active_names
        excluded = [str(i) for i in range(1, 16)
                    if not flags.get(f'has_scope_3_category_{i}', False)]
        print(f"[范围三过滤] 保留类别: {list(active_scope_3.keys())}")
        print(f"[范围三过滤] 排除类别: {excluded}")

    # 4. 渲染模板
    print("[步骤4] 渲染模板...")
    
    # 注册自定义 Jinja2 过滤器
    from jinja2 import Environment, ChainableUndefined
    from jinja2_filters import format_number, format_emission, register_filters_to_template

    env = Environment(undefined=ChainableUndefined)
    env.filters['cn_num'] = to_chinese_num
    env.filters['format_number'] = format_number
    env.filters['format_emission'] = format_emission

    print("[渲染] 已注册过滤器: cn_num, format_number, format_emission (未定义变量将显示为空)")

    today = date.today()
    render_context['posted_time'] = f"{today.year}年{today.month}月{today.day}日"
    print(f"[渲染] 生成日期（覆盖posted_time）: {render_context['posted_time']}")

    template.render(render_context, jinja_env=env)

    # 5. 保存报告
    print(f"\n[步骤5] 保存报告到: {output_path}")
    template.save(output_path)

    # 6. 统一公司简介和经营范围的段落格式
    print(f"\n[步骤6] 统一段落格式...")
    from docx import Document
    from docx.shared import Pt, Inches, Cm

    doc = Document(output_path)
    company_name = context.get('company_name', '')

    # 首行缩进：2个中文字符宽度（五号字 = 10.5pt，2字符 = 21pt ≈ 0.74cm）
    first_line_indent = Cm(0.74)

    def _strip_leading_spaces(para):
        """移除段落首 run 的前导空格，改用 first_line_indent 控制缩进"""
        if para.runs:
            first_run = para.runs[0]
            first_run.text = first_run.text.lstrip(' \t')

    def _remove_first_line_chars(para):
        """移除 w:firstLineChars 属性，避免其覆盖 w:firstLine 的缩进设置"""
        from docx.oxml.ns import qn
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:firstLineChars'), None)

    def _split_para_by_newlines(para):
        """将包含 \\n\\n 的段落拆分为多个段落，返回新创建的段落数"""
        text = para.text
        if '\n\n' not in text:
            return 0
        parts = [p.strip() for p in text.split('\n\n') if p.strip()]
        if len(parts) <= 1:
            return 0

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from copy import deepcopy

        # 保留原始 run 的格式（字体、大小等），取第一个 run 作为模板
        orig_runs = list(para.runs)
        ref_run = orig_runs[0] if orig_runs else None
        ref_rPr = deepcopy(ref_run._element.find(qn('w:rPr'))) if ref_run is not None and ref_run._element.find(qn('w:rPr')) is not None else None

        # 保存原始段落的行距设置（1.5 倍行距 = 360 twips）
        ref_pPr = para._element.find(qn('w:pPr'))
        ref_spacing = deepcopy(ref_pPr.find(qn('w:spacing'))) if ref_pPr is not None and ref_pPr.find(qn('w:spacing')) is not None else None

        # 第一个段落内容放入当前段落
        para.clear()
        run = para.add_run(parts[0])
        if ref_rPr is not None:
            run._element.insert(0, deepcopy(ref_rPr))

        # 在 body 中找到当前段落的位置，在其后插入新段落
        body = para._element.getparent()
        insert_idx = list(body).index(para._element)

        for part_text in parts[1:]:
            insert_idx += 1
            p_el = OxmlElement('w:p')
            # 段落属性：首行缩进 + 行距
            pPr = OxmlElement('w:pPr')
            ind_el = OxmlElement('w:ind')
            ind_el.set(qn('w:firstLine'), '420')  # 420 twips = 2 × 10.5pt
            ind_el.set(qn('w:left'), '0')
            pPr.append(ind_el)
            if ref_spacing is not None:
                pPr.append(deepcopy(ref_spacing))
            p_el.append(pPr)
            # 文本 run
            r_el = OxmlElement('w:r')
            if ref_rPr is not None:
                r_el.append(deepcopy(ref_rPr))
            t_el = OxmlElement('w:t')
            t_el.text = part_text
            t_el.set(qn('xml:space'), 'preserve')
            r_el.append(t_el)
            p_el.append(r_el)
            body.insert(insert_idx, p_el)

        return len(parts) - 1

    # 遍历所有段落，精确处理公司简介和经营范围
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查是否是公司简介/经营范围段落（以公司名开头，且非量化方法说明段落）
        is_profile = (company_name in text[:50] or '楚能新能源' in text[:50]) and len(text) > 100
        is_emission_item = any(kw in text for kw in ['活动数据AD', '排放因子EF', '量化模型'])
        if is_profile and not is_emission_item:
                n = _split_para_by_newlines(para)
                _strip_leading_spaces(para)
                para.paragraph_format.first_line_indent = first_line_indent
                para.paragraph_format.left_indent = 0
                # 移除模板遗留的 firstLineChars，避免覆盖 firstLine 的设置
                _remove_first_line_chars(para)
                print(f"  已处理公司简介段落（长度: {len(para.text)} 字符，拆分出 {n} 段）")

        # 检查是否是经营范围段落（包含经营范围特征）
        elif '经营范围' in text and len(text) > 100 and not is_emission_item:
            n = _split_para_by_newlines(para)
            _strip_leading_spaces(para)
            para.paragraph_format.first_line_indent = first_line_indent
            para.paragraph_format.left_indent = 0
            _remove_first_line_chars(para)
            print(f"  已处理经营范围段落（长度: {len(para.text)} 字符，拆分出 {n} 段）")

    # 7. 设置表2的列宽相等
    print(f"\n[步骤7] 设置表2列宽...")
    # 找到表2（范围二、三间接排放源表格）
    if len(doc.tables) >= 2:
        table2 = doc.tables[1]  # 第二个表格是表2
        # 设置所有列宽相等（使用 Inches 单位）
        equal_width = Inches(2.0)  # 每列2英寸
        for row in table2.rows:
            for cell in row.cells:
                # 设置单元格宽度
                cell.width = equal_width
        print("  表2列宽已设置为相等")

    doc.save(output_path)
    print("段落格式统一完成")

    # 7.5 插入范围三各类别材料列表副标题（标题和量化模型之间）
    print(f"\n[步骤7.5] 插入范围三各类别材料列表副标题...")
    _insert_category1_material_subtitle(doc, context)

    doc.save(output_path)
    print("材料列表副标题插入完成")

    # 7.6 统一全文英文括号为中文括号
    print(f"\n[步骤7.6] 统一中英文括号...")
    convert_parentheses_to_chinese(doc)

    doc.save(output_path)
    print("括号统一完成")

    # 7.7. 插入 Word TOC 域目录
    print(f"\n[步骤7.7] 插入 Word TOC 域目录...")
    insert_toc_field(doc)
    doc.save(output_path)

    # 8. 清理量化方法说明部分的过多空行
    print(f"\n[步骤8] 清理量化方法说明部分的空行...")
    clean_excessive_blank_lines(doc)

    doc.save(output_path)
    print("空行清理完成")

    # 9. 删除没有数据的类别表格（仅删除标题段落，保留表格结构）
    print(f"\n[步骤9] 删除没有数据的类别表格...")
    clean_empty_category_tables_v2(doc, context)

    doc.save(output_path)
    print("空类别表格清理完成")

    # 9.4. 修复范围三类别标题缺失类别名称的问题
    print(f"\n[步骤9.4] 修复范围三类别标题...")
    fix_scope3_category_headers(doc)

    doc.save(output_path)
    print("范围三类别标题修复完成")

    # 9.45. 添加范围三排除类别说明
    print(f"\n[步骤9.45] 添加范围三排除类别说明...")
    add_excluded_categories_statement(doc, context)

    # 10. 使用 XML vMerge 方法合并表格中的纵向单元格（针对表1和表2）
    print(f"\n[步骤10] 使用 XML vMerge 方法合并表格中的纵向单元格...")

    # 处理表1（doc.tables[0]）
    if len(doc.tables) >= 1:
        table1 = doc.tables[0]
        print(f"  处理表1（表格索引0）...")
        try:
            merge_vertical_cells(table1, 0)
        except Exception as e:
            print(f"  处理表1时出错: {e}")
            import traceback
            traceback.print_exc()

    # 处理表2（doc.tables[1]）
    if len(doc.tables) >= 2:
        table2 = doc.tables[1]
        print(f"  处理表2（表格索引1）...")
        try:
            merge_vertical_cells(table2, 0)
            _center_table_cells_horizontal(table2)
        except Exception as e:
            print(f"  处理表2时出错: {e}")
            import traceback
            traceback.print_exc()

    doc.save(output_path)
    print("表格纵向单元格合并完成（XML vMerge 方法）")

    # 10.5. 合并其他表格中的纵向单元格（XML方法）- 范围三类别表格
    print(f"\n[步骤10.5] 合并范围三类别表格的纵向单元格（XML方法）...")
    merge_other_tables_vertical_cells(doc, context)

    doc.save(output_path)
    print("范围三类别表格纵向单元格合并完成（XML方法）")

    # 10.6. 化学式下标转换
    print(f"\n[步骤10.6] 化学式数字下标转换...")
    apply_chemical_subscripts(doc)
    doc.save(output_path)
    print("化学式下标转换完成")

    # 10.7. 最终全局空行清理（简单可靠：删除连续空段落，最多保留1个）
    print(f"\n[步骤10.7] 最终全局空行清理...")
    final_removed = 0
    consecutive = 0
    to_remove = []
    ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            # 检查段落是否包含图片、图表等非文字内容
            drawings = para._element.findall('.//' + ns_w + 'drawing')
            if drawings:
                consecutive = 0  # 图片段落，不计为空段落
                continue
            consecutive += 1
            if consecutive > 1:
                to_remove.append(i)
        else:
            consecutive = 0
    for idx in reversed(to_remove):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
        final_removed += 1
    doc.save(output_path)
    print(f"  最终删除了 {final_removed} 个多余空行")

    # 10.8. 设置打开文档时自动更新域（TOC会提示更新，点一次"是"即可）
    print(f"\n[步骤10.8] 设置文档自动更新域...")
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    settings_element = doc.settings._element
    existing = settings_element.find(qn('w:updateFields'))
    if existing is not None:
        settings_element.remove(existing)
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_element.append(update_fields)
    doc.save(output_path)
    print("  已设置：打开文档时自动更新域（包括目录）")

    # 10.9. 使用 Word COM 自动化更新 TOC 域并获取正确页码
    print(f"\n[步骤10.9] 使用 Word COM 更新目录域（获取正确页码）...")
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            abs_path = os.path.abspath(output_path)
            wdoc = word.Documents.Open(abs_path)
            # 第一次更新：生成 TOC 条目
            wdoc.Fields.Update()
            # 强制重新分页（TOC 页码需要分页信息）
            wdoc.Repaginate()
            # 第二次更新：用正确的页码更新 TOC
            if wdoc.TablesOfContents.Count > 0:
                wdoc.TablesOfContents(1).Update()
            wdoc.Save()
            wdoc.Close()
            print("  Word COM 目录更新完成（页码已正确）")
        finally:
            word.Quit()
    except ImportError:
        print("  pywin32 未安装，跳过 COM 目录更新（页码可能不正确）")
    except Exception as e:
        print(f"  COM 目录更新失败: {e}")

    print("\n" + "=" * 50)
    print(f"报告生成成功: {output_path}")
    print("=" * 50)

    return output_path
if __name__ == "__main__":
    import sys

    # 检查命令行参数
    if len(sys.argv) > 1 and sys.argv[1] == '--generate':
        # 生成报告模式
        xlsx_path = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_DY_XLSX_NAME
        output_path = sys.argv[3] if len(sys.argv) > 3 else "carbon_report.docx"
        generate_report_from_xlsx(xlsx_path=xlsx_path, output_path=output_path)
    else:
        # 默认执行生成报告
        print("使用 'python main.py --generate' 生成报告")
        generate_report_from_xlsx(output_path="carbon_report.docx")
