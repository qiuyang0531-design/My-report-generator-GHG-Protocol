#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复template.docx中的Jinja2语法错误
"""

import docx
import re

def fix_template():
    """修复模板中的语法错误"""
    doc = docx.Document('template.docx')

    # 需要修复的错误标签
    fixes = {
        '{{ scope_of_ business }}': '{{ scope_of_business }}',
        # 可以添加其他需要修复的标签
    }

    # 修复段落中的标签
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for wrong_tag, correct_tag in fixes.items():
            if wrong_tag in text:
                paragraph.text = text.replace(wrong_tag, correct_tag)
                print(f"修复段落: {wrong_tag} -> {correct_tag}")

    # 修复表格中的标签
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for wrong_tag, correct_tag in fixes.items():
                        if wrong_tag in text:
                            paragraph.text = text.replace(wrong_tag, correct_tag)
                            print(f"修复表格: {wrong_tag} -> {correct_tag}")

    # 保存修复后的模板
    doc.save('template_fixed.docx')
    print("模板修复完成，保存为 template_fixed.docx")

if __name__ == "__main__":
    fix_template()