"""
Word 文档后处理模块
====================
负责对生成的 docx 报告进行格式修正、表格合并、化学式转换等后处理。
"""

import re
import traceback
from copy import deepcopy

from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches

from report_config import get_scope_3_category_name, get_scope_3_category_name_by_chinese_numeral


def find_table_by_content(doc, search_keywords):
    """
    根据表格内容动态查找表格索引

    Args:
        doc: Word文档对象
        search_keywords: 要搜索的关键词列表（表格中包含任一关键词即匹配）

    Returns:
        匹配的表格索引，如果未找到返回None
    """
    for idx, table in enumerate(doc.tables):
        # 获取表格中所有文本
        table_text = ""
        try:
            # 直接访问XML元素，避免python-docx的导航问题
            for row in table._element.tr_lst:
                for tc in row.tc_lst:
                    # 获取单元格文本
                    for p in tc.p_lst:
                        for r in p.r_lst:
                            for t in r.t_lst:
                                table_text += t.text + " "
        except Exception as e:
            # 如果遍历出错（如合并单元格），跳过该表格
            continue

        # 检查是否包含任一关键词
        for keyword in search_keywords:
            if keyword in table_text:
                return idx

    return None


# Unicode 下标数字映射
_SUB_DIGITS = str.maketrans('0123456789', '₀₁₂₃₄₅₆₇₈₉')

# 化学式精确替换列表：从长到短排列，避免部分匹配（CO2e 必须在 CO2 之前）
_CHEM_REPLACEMENTS = [
    ('CO2e', 'CO₂e'),
    ('CO2', 'CO₂'),
    ('CH4', 'CH₄'),
    ('N2O', 'N₂O'),
    ('SF6', 'SF₆'),
    ('NF3', 'NF₃'),
]


def apply_chemical_subscripts(doc):
    """
    将文档中所有化学式数字转为 Unicode 下标（CO2→CO₂, CH4→CH₄ 等）。
    遍历段落、表格、页眉页脚中的所有 run。
    """
    processed_count = 0

    def _needs_conversion(text):
        """检查文本是否含有需要转换的化学式（未下标的原始形式）"""
        for old, _new in _CHEM_REPLACEMENTS:
            if old in text:
                return True
        return False

    def _replace_in_text(text):
        """对文本中所有化学式数字做下标替换"""
        for old, new in _CHEM_REPLACEMENTS:
            text = text.replace(old, new)
        return text

    def _process_paragraph(para):
        nonlocal processed_count
        # 先尝试逐 run 替换
        any_changed = False
        for run in para.runs:
            old = run.text
            new = _replace_in_text(old)
            if new != old:
                run.text = new
                any_changed = True
                processed_count += 1

        # 若段落全文仍有未转换的化学式（跨 run 拆分导致），合并所有 run
        if _needs_conversion(para.text):
            full = _replace_in_text(para.text)
            if para.runs:
                para.runs[0].text = full
                for r in para.runs[1:]:
                    r.text = ''
                processed_count += 1

    def _process_document(d):
        for para in d.paragraphs:
            _process_paragraph(para)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_paragraph(para)
        for section in d.sections:
            header = section.header
            if header:
                for para in header.paragraphs:
                    _process_paragraph(para)
            footer = section.footer
            if footer:
                for para in footer.paragraphs:
                    _process_paragraph(para)

    _process_document(doc)
    print(f"  化学式下标转换: {processed_count} 处")


def _insert_category1_material_subtitle(doc, context):
    """在范围三各类别标题后插入材料列表副标题段落"""
    scope_3_names = context.get('scope_3_category_names', {})
    if not scope_3_names:
        print("  未找到 scope_3_category_names，跳过")
        return

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    inserted = 0
    for cat_key, cat_name in scope_3_names.items():
        # 提取类别编号
        try:
            cat_num = int(cat_key.split('_')[1])
        except (IndexError, ValueError):
            continue

        subtitle_key = f'scope3_category{cat_num}_material_subtitle'
        subtitle = context.get(subtitle_key, '')
        if not subtitle:
            continue

        # 查找标题段落：匹配中文编号 + 类别名称
        target_para = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r'（[一二三四五六七八九十]+）', text) and cat_name in text:
                target_para = para
                break

        if target_para is None:
            continue

        target_el = target_para._element

        # 创建新段落
        new_p = OxmlElement('w:p')

        # 复制段落属性（pPr），移除编号属性避免被自动编号
        target_pPr = target_el.find(qn('w:pPr'))
        if target_pPr is not None:
            new_pPr = deepcopy(target_pPr)
            numPr = new_pPr.find(qn('w:numPr'))
            if numPr is not None:
                new_pPr.remove(numPr)
            ind = new_pPr.find(qn('w:ind'))
            if ind is not None:
                ind.attrib.pop(qn('w:firstLine'), None)
                ind.attrib.pop(qn('w:hanging'), None)
            new_p.append(new_pPr)

        # 创建 run 并设置文本
        r_el = OxmlElement('w:r')
        target_runs = target_el.findall(qn('w:r'))
        if target_runs:
            target_rPr = target_runs[0].find(qn('w:rPr'))
            if target_rPr is not None:
                r_el.append(deepcopy(target_rPr))

        t_el = OxmlElement('w:t')
        t_el.text = subtitle
        t_el.set(qn('xml:space'), 'preserve')
        r_el.append(t_el)
        new_p.append(r_el)

        # 插入到标题段落后
        parent = target_el.getparent()
        target_idx = list(parent).index(target_el)
        parent.insert(target_idx + 1, new_p)
        inserted += 1

    if inserted:
        print(f"  已插入 {inserted} 个范围三类别材料列表副标题")
    else:
        print("  没有需要插入的材料列表副标题")


def convert_parentheses_to_chinese(doc):
    """
    将全文中的英文括号 () 统一替换为中文括号 （）
    覆盖段落和表格中的所有文本 run。
    """
    print("  正在统一中英文括号...")
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if '(' in run.text or ')' in run.text:
                run.text = run.text.replace('(', '（').replace(')', '）')
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if '(' in run.text or ')' in run.text:
                            run.text = run.text.replace('(', '（').replace(')', '）')
                            count += 1
    print(f"  统一了 {count} 处括号")


def clean_excessive_blank_lines(doc):
    """
    清理量化方法说明部分过多的空行
    策略：每个排放源项的描述后保留一个空行，删除描述前的空行
    """
    import re
    print("  正在清理量化方法说明部分的空行...")

    # 找到量化方法说明章节
    start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if '量化方法说明' in para.text:
            start_idx = i
            break

    if not start_idx:
        print("  未找到量化方法说明章节")
        return

    # 找到该部分的结束位置
    end_idx = start_idx + 1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        # 找到下一个主要章节
        if text and ('四、' in text or '参考文献' in text or '附录' in text):
            end_idx = i
            break
        if i > start_idx + 200:
            end_idx = i
            break

    print(f"  量化方法说明部分: {start_idx} 到 {end_idx}")

    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    def _is_visually_empty(para):
        """判断段落是否视觉上为空（无可见文字，可能含 <w:br/> 标签）"""
        text = para.text.strip()
        if text:
            return False
        runs = para._element.findall(ns + 'r')
        if not runs:
            return True
        for run in runs:
            t_elems = run.findall(ns + 't')
            run_text = ''.join(t.text or '' for t in t_elems).strip()
            if run_text:
                return False
        return True

    # 1. 删除视觉空段落，保留最多1个空行
    indices_to_remove = set()
    consecutive_empty = 0

    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break
        if _is_visually_empty(doc.paragraphs[i]):
            consecutive_empty += 1
            if consecutive_empty > 1:
                indices_to_remove.add(i)
        else:
            consecutive_empty = 0

    # 删除连续空段落（保留最多1个空行）
    removed_count = 0
    for idx in sorted(indices_to_remove, reverse=True):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            parent = para._element.getparent()
            parent.remove(para._element)
            removed_count += 1

    print(f"  删除了 {removed_count} 个多余空行")

    # 5. 清理段落内部所有的 <w:br/> 空 run（不仅尾随，包括全部）
    br_cleaned = 0
    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        runs = para._element.findall(ns + 'r')
        if not runs:
            continue
        runs_to_remove = []
        for run in runs:
            brs = run.findall(ns + 'br')
            if not brs:
                continue
            t_elems = run.findall(ns + 't')
            t_text = ''.join(t.text or '' for t in t_elems).strip()
            if not t_text:
                runs_to_remove.append(run)
        for run in runs_to_remove:
            para._element.remove(run)
            br_cleaned += 1

    if br_cleaned:
        print(f"  清理了 {br_cleaned} 个段落内部的残留 <w:br/> 空行")

def insert_toc_field(doc):
    """在"概述"段落前插入 Word TOC 域目录

    1. 给所有 Style2 段落设置大纲级别 1（排除"目录"标题自身）
    2. 插入 TOC \\o "1-1" 域（基于大纲级别，比 \\t 样式名匹配更可靠）
    3. 设置 updateFields，打开文档即可更新（点一次"是"）
    """
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml, OxmlElement

    # ---- Step 1: 给 Style2 段落设置大纲级别 ----
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    outline_count = 0
    for para in doc.paragraphs:
        if para.style.name == 'Style2':
            text = para.text.strip()
            if text == '目录':
                continue  # 目录标题自身不入目录
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._element.insert(0, pPr)
            # 移除旧的大纲级别
            old_lvl = pPr.find(qn('w:outlineLvl'))
            if old_lvl is not None:
                pPr.remove(old_lvl)
            # 设置大纲级别 0 = Level 1（对应 TOC \\o "1-1"）
            outline_lvl = OxmlElement('w:outlineLvl')
            outline_lvl.set(qn('w:val'), '0')
            pPr.append(outline_lvl)
            outline_count += 1
    print(f"  已为 {outline_count} 个 Style2 段落设置大纲级别 1")

    # ---- Step 2: 查找"概述"段落 ----
    target = None
    for para in doc.paragraphs:
        if para.text.strip() == '概述' and para.style.name == 'Style2':
            target = para._element
            break

    if target is None:
        print("  未找到'概述'段落，跳过目录插入")
        return

    # ---- Step 3: 目录标题段落（Normal 样式 + 加粗居中 + 段前新页，不入目录） ----
    toc_title = parse_xml(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:pPr>'
        '<w:pageBreakBefore/>'
        '<w:keepNext/>'
        '<w:jc w:val="center"/>'
        '</w:pPr>'
        '<w:r>'
        '<w:rPr><w:b/><w:sz w:val="32"/></w:rPr>'
        '<w:t xml:space="preserve">目录</w:t>'
        '</w:r>'
        '</w:p>'
    )

    # ---- Step 4: TOC 域段落（\\o "1-1" 收集大纲级别 1 的段落） ----
    toc_field = parse_xml(
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r>'
        '<w:fldChar w:fldCharType="begin"/>'
        '</w:r>'
        '<w:r>'
        '<w:instrText xml:space="preserve"> TOC \\o "1-1" \\h </w:instrText>'
        '</w:r>'
        '<w:r>'
        '<w:fldChar w:fldCharType="separate"/>'
        '</w:r>'
        '<w:r>'
        '<w:rPr><w:b/><w:color w:val="FF0000"/></w:rPr>'
        '<w:t>（打开文档时点"是"更新目录）</w:t>'
        '</w:r>'
        '<w:r>'
        '<w:fldChar w:fldCharType="end"/>'
        '</w:r>'
        '</w:p>'
    )

    # ---- Step 5: 给"概述"段落加 pageBreakBefore，确保正文从新页开始 ----
    target_pPr = target.find(qn('w:pPr'))
    if target_pPr is None:
        target_pPr = OxmlElement('w:pPr')
        target.insert(0, target_pPr)
    page_break = OxmlElement('w:pageBreakBefore')
    target_pPr.append(page_break)

    # 插入顺序：toc_title → toc_field → 概述
    target.addprevious(toc_field)       # TOC 域紧贴概述
    toc_field.addprevious(toc_title)    # "目录"标题在 TOC 域前

    print('  TOC 域目录已插入（大纲级别模式，目录另起一页，正文从新页开始）')



def fix_scope3_category_headers(doc):
    """
    修复第四章量化说明中范围三类别标题缺失类别名称的问题

    修复内容：
    - 将 "（一）" 替换为 "（一）购买的商品和服务"
    - 将 "（二）" 替换为 "（二）资本货物"
    - 其他类别的类似标题
    """
    print("  正在修复范围三类别标题...")

    # 找到"范围三：其他间接温室气体排放"的位置
    scope3_section_start = None
    paragraphs_list = list(doc.paragraphs)

    for i, para in enumerate(paragraphs_list):
        text = para.text.strip()
        if '范围三' in text and '其他间接温室气体排放' in text:
            scope3_section_start = i
            break

    if scope3_section_start is None:
        print("    未找到范围三章节，跳过修复")
        return

    print(f"    找到范围三章节在段落 {scope3_section_start}")

    # 确保范围三分组标题不带编号（如已存在"（三）"则去除）
    for i in range(max(0, scope3_section_start - 10), min(len(paragraphs_list), scope3_section_start + 5)):
        if i < len(paragraphs_list):
            text = paragraphs_list[i].text.strip()
            if '（三）' in text and ('范围三' in text or '其他类别' in text):
                para = paragraphs_list[i]
                for run in para.runs:
                    if '（三）' in run.text:
                        run.text = run.text.replace('（三）', '')
                        break
                print(f"    已移除段落 {i} 的 '（三）' 前缀")
                break

    # 在范围三章节内查找只有编号没有类别名称的标题
    fixed_count = 0
    title_pattern = re.compile(r'^（[一二三四五六七八九十]+）')
    for i in range(scope3_section_start + 1, len(paragraphs_list)):
        text = paragraphs_list[i].text.strip()

        # 如果到达下一个范围章节，停止处理
        if text and ('第四章' in text or '参考文献' in text or '附录' in text):
            break

        # 情况1：纯编号标题（只有"（X）"而没有其他内容）
        if get_scope_3_category_name_by_chinese_numeral(text) != text:
            if i + 1 < len(paragraphs_list):
                next_text = paragraphs_list[i + 1].text.strip()
                if next_text.startswith('（1') or next_text.startswith('（2'):
                    full_title = f"{text}{get_scope_3_category_name_by_chinese_numeral(text)}排放"
                    para = paragraphs_list[i]
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = full_title
                    else:
                        para.add_run(full_title)
                    fixed_count += 1
                    print(f"    修复段落{i}: '{text}' -> '{full_title}'")
            continue

        # 情况2：已有类别名称但缺少"排放"后缀
        #    匹配所有 （中文数字）... 开头的标题，不含"排放"则追加
        #    注意：文本可能跨多个 run，不能直接 replace，需清空后重写
        if title_pattern.match(text) and '排放' not in text:
            full_title = f"{text}排放"
            para = paragraphs_list[i]
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = full_title
            else:
                para.add_run(full_title)
            fixed_count += 1
            print(f"    修复段落{i}: '{text}' -> '{full_title}'")

    print(f"    已修复 {fixed_count} 个范围三类别标题")


def add_excluded_categories_statement(doc, context):
    """
    在报告合规声明区域添加范围三排除类别说明

    列出本次盘查范围内不进行量化的范围三类别（8, 13, 14, 15），
    并说明原因（数据不具备重要性）。
    """
    print("  正在添加范围三排除类别说明...")

    flags = context.get('flags', {})
    # 收集所有排除的类别（flag=False 的类别）
    excluded_nums = []
    for i in range(1, 16):
        flag_key = f'has_scope_3_category_{i}'
        if not flags.get(flag_key, False):
            excluded_nums.append(i)

    if not excluded_nums:
        print("    所有范围三类别均有排放数据，无需添加排除说明")
        return

    # 获取类别名称
    scope_3_names = context.get('scope_3_category_names', {})
    # scope_3_category_names 可能使用整数键或字符串键
    excluded_names = []
    for num in excluded_nums:
        name = scope_3_names.get(f'category_{num}') or scope_3_names.get(num) or f'类别{num}'
        excluded_names.append(f'类别{num}（{name}）')

    # 生成排除说明文字
    names_str = '、'.join([scope_3_names.get(f'category_{n}') or scope_3_names.get(n) or f'类别{n}' for n in excluded_nums])
    statement = f"本次盘查范围内，无{names_str}相关排放。"

    # 找到最后一个合规声明段落（无温室气体储存）作为插入参考点
    insert_after = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if '无温室气体储存' in text or 'GHG Storage' in text:
            insert_after = para
            break

    if insert_after is None:
        # Fallback: 找无生物质燃烧排放之后的段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if '无生物质燃烧排放' in text:
                insert_after = para
                # 继续找最后一个声明
                for p2 in doc.paragraphs:
                    if '无温室气体储存' in p2.text or 'GHG Storage' in p2.text:
                        insert_after = p2
                break

    if insert_after is None:
        print("    未找到合规声明段落，跳过")
        return

    # 在合规声明段落后插入排除说明
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from copy import deepcopy

    # 克隆参考段落的 XML 元素以保持样式
    ref_element = insert_after._element
    new_p = OxmlElement('w:p')

    # 复制段落属性（样式、缩进等）
    ref_pPr = ref_element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = deepcopy(ref_pPr)
        new_p.append(new_pPr)

    # 创建新的 run 并设置文本
    new_r = OxmlElement('w:r')
    # 复制 run 属性（字体等）从参考段落的第一个 run
    ref_rPr = ref_element.find(qn('w:r') + '/' + qn('w:rPr'))
    if ref_rPr is not None:
        new_rPr = deepcopy(ref_rPr)
        new_r.append(new_rPr)

    new_t = OxmlElement('w:t')
    new_t.text = statement
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)

    # 在参考段落后插入
    ref_element.addnext(new_p)
    print(f"    已添加排除类别说明: {statement[:80]}...")


def clean_empty_category_tables_v2(doc, context):
    """
    删除没有数据的类别的标题段落和单位段落，但保留表格结构

    新版本：不删除表格本身，避免表格索引偏移问题
    只删除与空类别相关的标题段落和单位段落
    """
    TOTAL_SCOPE3_CATEGORIES = 15

    # 检查哪些类别完全没有任何数据（既没有ef_items，也没有detail_items，也没有emissions）
    empty_categories = []
    # 同时检查哪些类别的排放因子表只有表头（没有数据行）
    empty_ef_table_categories = []

    for i in range(1, TOTAL_SCOPE3_CATEGORIES + 1):
        ef_items = context.get(f'cat{i}_ef_items', [])
        detail_items = context.get(f'scope3_category{i}', [])
        emission_value = context.get(f'scope_3_category_{i}_emissions', 0)

        has_ef_items = ef_items and len(ef_items) > 0
        has_detail_items = detail_items and len(detail_items) > 0
        has_emissions = emission_value and emission_value > 0

        # 如果没有任何数据，则标记为完全空类别
        if not (has_ef_items or has_detail_items or has_emissions):
            empty_categories.append(i)

        # 如果排放因子表没有数据（ef_items为空），标记为需要删除排放因子表标题
        if not has_ef_items:
            empty_ef_table_categories.append(i)


    if not empty_categories and not empty_ef_table_categories:
        print("  所有类别都有数据，无需删除空类别标题")
        return

    print(f"  完全空类别（无任何数据）: {empty_categories}")
    print(f"  排放因子表为空的类别: {[c for c in empty_ef_table_categories if c not in empty_categories]}")

    # 收集要删除的段落
    all_paragraphs_to_remove = []

    # 步骤1：查找并删除完全空类别的标题段落（只删除真正空的类别，不删除只有EF表为空但有其他数据的类别）
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 只检查完全空类别（不包含只有EF表为空的类别）
        is_empty_category_para = False

        for cat_num in empty_categories:  # 只使用 empty_categories，不使用 empty_ef_table_categories
            category_name = get_scope_3_category_name(cat_num)
            is_target_category = (
                f'范围三 类别{cat_num}' in text or
                f'范围三类别{cat_num}' in text or
                f'类别{cat_num} ' in text or  # 类别X后面有空格，避免匹配"类别10"到"类别1"
                (category_name and category_name in text)
            )

            if is_target_category:
                is_empty_category_para = True
                break

        if is_empty_category_para:
            all_paragraphs_to_remove.append(para)

    # 步骤2：删除所有标记的段落
    deleted_count = 0
    for para in all_paragraphs_to_remove:
        try:
            para_element = para._element
            parent = para_element.getparent()
            if parent is not None:
                parent.remove(para_element)
                deleted_count += 1
        except Exception as e:
            print(f"  删除段落时出错: {e}")

    print(f"  已删除 {deleted_count} 个空类别相关段落")

    # 步骤2.5：删除孤立的单位描述段落（类别表格被删除后留下的单位描述）
    # 这些段落通常只包含"单位：吨CO2e"或类似内容

    # 删除孤立的单位列/段落
    orphan_unit_paras = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # 检查是否是单位描述段落
        if text in ['单位：吨CO2e', '单位: 吨CO2e', '单位：tCO2e', '单位: tCO2e', '单位：吨二氧化碳当量', '单位: 吨二氧化碳当量']:
            # 检查这个段落是否在已删除类别的位置附近
            # 通过检查前面是否有已删除类别的表格来判断
            orphan_unit_paras.append(para)
        # 也处理只有"单位"或"Unit"的情况
        elif text in ['单位', 'Unit', '单位：', '单位:']:
            orphan_unit_paras.append(para)

    # 删除孤立的单位段落
    for para in orphan_unit_paras:
        try:
            para_element = para._element
            parent = para_element.getparent()
            if parent is not None:
                parent.remove(para_element)
                deleted_count += 1
        except Exception as e:
            print(f"  删除单位段落时出错: {e}")

    if orphan_unit_paras:
        print(f"  已删除 {len(orphan_unit_paras)} 个孤立单位描述段落")

    # 步骤3：删除空类别相关的表格
    tables_to_remove = []

    # 定义表头关键词列表
    header_keywords = ['编号', 'GHG排放类别', '排放源', 'Activity name',
                     'Geography', 'CO2', 'CH4', 'N2O', '单位', '引用源',
                     '排放因子', '缺省', '基于热值']

    # 动态查找每个类别对应的表格（区分库存表和EF表）
    category_inventory_table_indices = {}  # 库存表：包含3.10.1这类详细数据
    category_ef_table_indices = {}  # EF表：包含排放因子数据
    for cat_num in range(1, 16):
        # 在所有表格中查找属于该类别的表格（从表格4开始，包含范围三类别汇总表）
        for table_idx in range(4, min(50, len(doc.tables))):
            if table_idx in category_inventory_table_indices.values() or table_idx in category_ef_table_indices.values():
                continue  # 已经被其他类别占用

            table = doc.tables[table_idx]
            row_count = len(table.rows)

            # 检查表格是否属于该类别
            is_category_table = False
            is_inventory_table = False  # 是否是库存表（包含3.10.1这类编号）
            is_ef_table = False  # 是否是EF表（包含排放因子相关列）

            for row_idx in range(min(6, row_count)):
                row = table.rows[row_idx]
                for cell in row.cells:
                    text = cell.text.strip()
                    if f'范围三 类别{cat_num}' in text or f'范围三类别{cat_num}' in text:
                        is_category_table = True
                    # 检查是否包含详细数据编号（如3.10.1）- 库存表判断
                    if f'3.{cat_num}.' in text or f'3.{cat_num} ' in text:
                        is_inventory_table = True
                    # 检查是否是EF表（包含排放因子相关关键词）- 只有在不是库存表的情况下才判断
                    if not is_inventory_table and any(keyword in text for keyword in ['排放因子', '缺省', '引用源', 'Activity name', 'Geography']):
                        if f'类别{cat_num}' in text or f'范围{cat_num}' in text:
                            is_ef_table = True
                if is_category_table:
                    break

            if is_category_table:
                # 根据表格内容判断类型 - 确保库存表优先级最高
                if is_inventory_table:
                    category_inventory_table_indices[cat_num] = table_idx
                elif is_ef_table:
                    category_ef_table_indices[cat_num] = table_idx
                else:
                    # 如果无法确定类型，先检查是否有类似3.X.Y的编号模式，再决定类型
                    has_detail_number_pattern = False
                    for row_idx in range(min(6, row_count)):
                        row = table.rows[row_idx]
                        for cell in row.cells:
                            text = cell.text.strip()
                            if re.search(r'3\.\d+\.\d+', text):
                                has_detail_number_pattern = True
                                break
                        if has_detail_number_pattern:
                            break

                    if has_detail_number_pattern:
                        category_inventory_table_indices[cat_num] = table_idx
                    else:
                        category_ef_table_indices[cat_num] = table_idx
                
                # 注意：不要 break，继续寻找该类别的其他表格（一个类别可能有 EF 表和 库存表）
                # continue

    # 检查哪些类别的表格需要删除
    # 1. 完全空类别（没有任何数据）- 删除所有相关表格
    for cat_num in empty_categories:
        # 删除库存表
        if cat_num in category_inventory_table_indices:
            table_idx = category_inventory_table_indices[cat_num]
            if table_idx < len(doc.tables):
                tables_to_remove.append(table_idx)
                print(f"  标记删除类别{cat_num}的库存表（完全空类别）: 索引{table_idx}")
        # 删除EF表
        if cat_num in category_ef_table_indices:
            table_idx = category_ef_table_indices[cat_num]
            if table_idx < len(doc.tables) and table_idx not in tables_to_remove:
                tables_to_remove.append(table_idx)
                print(f"  标记删除类别{cat_num}的EF表（完全空类别）: 索引{table_idx}")

    # 2. 只有EF表为空的类别（有详细数据或排放量）- 只删除EF表，保留库存表
    for cat_num in empty_ef_table_categories:
        if cat_num not in empty_categories:  # 跳过完全空类别（已处理）
            # 只删除EF表，不删除库存表
            if cat_num in category_ef_table_indices:
                table_idx = category_ef_table_indices[cat_num]
                if table_idx < len(doc.tables):
                    tables_to_remove.append(table_idx)
                    print(f"  标记删除类别{cat_num}的EF表（EF表为空但有库存数据）: 索引{table_idx}")

    # 同时检查所有范围三类别表格，删除明显的空表格
    # 搜索从表格4开始的所有表格（第四章范围三表格从表格4开始）
    all_category_tables = set(category_inventory_table_indices.values()) | set(category_ef_table_indices.values())
    for table_idx in range(4, len(doc.tables)):  # 搜索所有范围三表格
        if table_idx in tables_to_remove or table_idx in all_category_tables:
            continue  # 已经处理过

        table = doc.tables[table_idx]
        row_count = len(table.rows)

        # 对于3行或更少的表格，检查是否真的为空
        if row_count <= 3:
            # 检查表格是否是范围三类别表格
            is_scope3_table = False
            table_contains_unit_only = False  # 新增：标记表格只包含单位行

            for row in table.rows[:3]:
                for cell in row.cells:
                    text = cell.text.strip()
                    if 'GHG排放类别' in text or '范围三' in text:
                        is_scope3_table = True
                    # 检查是否只包含"单位：吨CO2e"这类内容
                    if text in ['单位：吨CO2e', '单位: 吨CO2e', '单位：tCO2e', '单位: tCO2e']:
                        table_contains_unit_only = True

            if is_scope3_table:
                # 对于3行的表格，检查第3行是否有数据（编号>0或非空内容）
                has_data = False
                if row_count >= 3:
                    for cell in table.rows[2].cells:
                        text = cell.text.strip()
                        # 检查是否有数字编号
                        try:
                            if float(text) > 0:
                                has_data = True
                                break
                        except (ValueError, TypeError):
                            # 检查是否有非空且有意义的文本（超过3个字符且不是表头关键词或单位说明）
                            if (len(text) > 3 and
                                text not in ['编号', 'GHG排放类别', '排放源', 'Activity name',
                                            'Geography', 'CO2', 'CH4', 'N2O', '单位', '引用源',
                                            '缺省排放因子', '单位：吨CO2e', '单位: 吨CO2e',
                                            '单位：tCO2e', '单位: tCO2e']):
                                has_data = True
                                break

                # 如果没有数据，或者表格只包含单位行，标记为删除
                if not has_data or table_contains_unit_only:
                    tables_to_remove.append(table_idx)
                    print(f"  标记删除明显空表格: 索引{table_idx}（范围三表格，只有表头/单位行，行数={row_count}）")
            continue

        # 对于4-6行的表格，检查是否有数据行
        if 4 <= row_count <= 6:
            # 跳过非排放数据的表格（如 GWP 表、活动数据汇总表等）
            header_text = ''
            if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
                header_text = table.rows[0].cells[0].text.strip()
            skip_keywords = ['GWP', '温室气体', '全球暖化', '活动数据汇总', '排放因子汇总']
            if any(kw in header_text for kw in skip_keywords):
                continue

            has_data = False
            # 跳过最后可能的"单位"行
            data_row_end = row_count - 1
            for row_idx in range(2, data_row_end):  # 从第2行开始检查到倒数第二行
                row = table.rows[row_idx]
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    # 跳过单位行
                    if first_cell_text in ['单位：吨CO2e', '单位: 吨CO2e', '单位：tCO2e', '单位: tCO2e', '单位', 'Unit']:
                        continue
                    try:
                        num_val = float(first_cell_text)
                        if num_val > 0:  # 有编号，说明有数据行
                            has_data = True
                            break
                    except (ValueError, TypeError):
                        # 检查是否有非空且有意义的文本（>=3 避免短字符串如 CH4 被误判为空）
                        if len(first_cell_text) >= 3:
                            has_data = True
                            break

            if not has_data:
                tables_to_remove.append(table_idx)
                print(f"  标记删除空表格: 索引{table_idx}（无数据行，行数={row_count}）")
            continue

        # 对于7行以上的表格，检查是否只有表头和单位行（没有实际数据）
        if row_count >= 7:
            has_real_data = False
            # 从第2行开始检查，跳过可能的单位行
            for row_idx in range(2, row_count - 1):
                if row_idx >= len(table.rows):
                    break
                row = table.rows[row_idx]
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    # 跳过单位行
                    if first_cell_text in ['单位：吨CO2e', '单位: 吨CO2e', '单位：tCO2e', '单位: tCO2e', '单位', 'Unit']:
                        continue
                    try:
                        num_val = float(first_cell_text)
                        if num_val > 0:  # 有编号，说明有数据行
                            has_real_data = True
                            break
                    except (ValueError, TypeError):
                        # 检查是否有非空且有意义的文本
                        if len(first_cell_text) > 3 and first_cell_text not in ['编号', 'GHG排放类别', '排放源', 'Activity name', 'Geography', 'CO2', 'CH4', 'N2O', '单位', '引用源']:
                            has_real_data = True
                            break

            if not has_real_data:
                tables_to_remove.append(table_idx)
                print(f"  标记删除空表格: 索引{table_idx}（只有表头/单位行，行数={row_count}）")

    # 从后往前删除表格
    deleted_table_count = 0
    deleted_title_count = 0
    for table_idx in sorted(tables_to_remove, reverse=True):
        if table_idx < len(doc.tables):
            # 先查找并删除表格前面的标题段落
            # EF表标题通常在表格前的一个段落，包含"排放因子表"字样
            table = doc.tables[table_idx]
            table_element = table._element

            # 获取表格在文档中的位置
            # 通过遍历所有表格元素来找到当前表格的索引
            all_table_elements = [t._element for t in doc.tables]
            current_position = all_table_elements.index(table_element)

            # 查找表格前可能存在的标题段落
            # 遍历所有段落，找到在表格之前的最后一个段落
            para_before_table = None
            for para in doc.paragraphs:
                para_element = para._element
                para_position = -1
                try:
                    # 获取段落位置
                    body = table_element.getparent().getparent()
                    if hasattr(body, 'index'):
                        para_position = body.index(para_element)
                except:
                    pass

                # 简单方法：通过检查段落内容是否匹配EF表标题模式
                para_text = para.text.strip()
                if ('排放因子表' in para_text or 'EF表' in para_text) and para_text:
                    # 检查是否是要删除类别的标题
                    for deleted_cat in empty_categories + empty_ef_table_categories:
                        if f'类别{deleted_cat}' in para_text or f'类别{deleted_cat}排放因子表' in para_text:
                            # 检查这个段落是否在表格附近（通过检查它后面的元素是否是表格）
                            # 获取段落的下一个兄弟元素
                            next_sibling = para_element.getnext()
                            if next_sibling is not None:
                                # 检查下一个元素是否是当前表格或指向当前表格
                                try:
                                    next_table = next_sibling
                                    # 检查是否是表格元素
                                    if next_table.tag.endswith('}tbl'):
                                        # 找到了标题段落
                                        para_before_table = para
                                        break
                                except:
                                    pass
                    if para_before_table:
                        break

            # 删除标题段落
            if para_before_table:
                try:
                    para_element = para_before_table._element
                    parent = para_element.getparent()
                    if parent is not None:
                        parent.remove(para_element)
                        deleted_title_count += 1
                        print(f"  删除表格{table_idx}前的EF表标题段落")
                except Exception as e:
                    print(f"  删除标题段落时出错: {e}")

            # 删除表格
            table_element = table._element
            table_element.getparent().remove(table_element)
            deleted_table_count += 1

    print(f"  已删除 {deleted_table_count} 个空类别表格和 {deleted_title_count} 个EF表标题段落")


def merge_vertical_cells(table, col_idx):
    """
    纵向合并表格指定列中内容相同的相邻单元格

    使用底层 XML 操作，正确设置 vMerge 属性来实现合并。
    直接访问 XML 元素，避免 python-docx 的导航问题。

    Args:
        table: python-docx 表格对象
        col_idx: 要处理的列索引（从0开始）

    Returns:
        合并的单元格数量
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if col_idx >= len(table.columns):
        print(f"  警告：列索引 {col_idx} 超出表格范围（表格共 {len(table.columns)} 列）")
        return 0

    merged_count = 0
    rows_count = len(table.rows)

    # 跳过表头行，从第1行开始处理（索引1）
    start_row = 1

    # 第一步：直接从 XML 读取所有行的文本内容
    print(f"  开始识别合并组（总行数：{rows_count}，从行{start_row}开始）...")

    # 获取表格的 XML 元素
    table_element = table._element
    tr_lst = table_element.tr_lst

    # 存储每行的文本内容
    row_texts = []

    for row_idx in range(len(tr_lst)):
        if row_idx < start_row:
            # 跳过表头行
            row_texts.append("")
            continue

        try:
            tr = tr_lst[row_idx]
            tc_lst = tr.tc_lst

            if col_idx >= len(tc_lst):
                print(f"  警告：行{row_idx}的列索引{col_idx}超出范围")
                row_texts.append("")
                continue

            tc = tc_lst[col_idx]

            # 获取单元格文本
            cell_text = ""
            for p in tc.p_lst:
                for r in r_lst if (r_lst := p.r_lst) else []:
                    for t in t_lst if (t_lst := r.t_lst) else []:
                        cell_text += t.text

            cell_text = cell_text.strip()
            row_texts.append(cell_text)
        except Exception as e:
            print(f"  警告：读取行{row_idx}时出错: {e}")
            row_texts.append("")

    # 第二步：识别需要合并的单元格组
    merge_groups = []
    group_start = start_row
    current_value = None

    for row_idx in range(start_row, rows_count):
        cell_text = row_texts[row_idx] if row_idx < len(row_texts) else ""

        # 初始化当前值（第一行）
        if current_value is None:
            current_value = cell_text
            group_start = row_idx
            text_preview = cell_text[:40] if cell_text else ""
            print(f"  行{row_idx}: 初始化，内容=\"{text_preview}\"")
            continue

        # 如果当前单元格内容与前一单元格相同，继续累积
        if cell_text == current_value and cell_text != "":
            continue

        # 内容不同或遇到空值，记录之前的合并组
        if row_idx - 1 > group_start:
            merge_groups.append((group_start, row_idx - 1))
            value_preview = current_value[:40] if current_value else ""
            print(f"  识别合并组: 行{group_start}-{row_idx-1} (共{row_idx-1-group_start}行), 内容=\"{value_preview}\"")

        # 开始新的合并组
        current_value = cell_text
        group_start = row_idx
        if cell_text:
            text_preview = cell_text[:40] if cell_text else ""
            print(f"  行{row_idx}: 新组开始，内容=\"{text_preview}\"")

    # 处理最后一组
    if rows_count - 1 > group_start:
        merge_groups.append((group_start, rows_count - 1))
        value_preview = current_value[:40] if current_value else ""
        print(f"  识别合并组: 行{group_start}-{rows_count-1} (共{rows_count-1-group_start}行), 内容=\"{value_preview}\"")

    print(f"  总共识别到 {len(merge_groups)} 个合并组")

    # 第三步：执行合并操作（直接操作 XML）
    print(f"  开始执行合并操作，共 {len(merge_groups)} 个合并组...")

    for group_start, group_end in merge_groups:
        if group_end <= group_start:
            print(f"  跳过无效合并组: 行{group_start}-{group_end}")
            continue

        print(f"  处理合并组: 行{group_start}-{group_end} (共{group_end-group_start}行)")

        try:
            # 获取顶部单元格的 XML 元素
            tr_start = tr_lst[group_start]
            tc_lst_start = tr_start.tc_lst

            if col_idx >= len(tc_lst_start):
                print(f"    警告：列索引{col_idx}超出范围")
                continue

            tc_start = tc_lst_start[col_idx]

            # 保存顶部单元格的文本内容
            merged_text = row_texts[group_start] if group_start < len(row_texts) else ""
            text_preview = merged_text[:40] if merged_text else ""
            print(f"    顶部单元格（行{group_start}）内容: \"{text_preview}\"")

            # 设置顶部单元格的 vMerge 属性为 "restart"
            tc_pr = tc_start.get_or_add_tcPr()
            for old_vmerge in tc_pr.findall(qn('w:vMerge')):
                tc_pr.remove(old_vmerge)
            v_merge = OxmlElement('w:vMerge')
            v_merge.set(qn('w:val'), 'restart')
            tc_pr.append(v_merge)

            # 验证vMerge属性是否设置成功
            verify_vm = tc_pr.find(qn('w:vMerge'))
            verify_val = verify_vm.get(qn('w:val')) if verify_vm is not None else 'not found'
            print(f"    设置行{group_start} vMerge=\"restart\" (验证: {verify_val})")

            # 处理其他单元格（设置为 "continue"）
            for row_idx in range(group_start + 1, group_end + 1):
                if row_idx >= len(tr_lst):
                    break

                try:
                    tr = tr_lst[row_idx]
                    tc_lst = tr.tc_lst

                    if col_idx >= len(tc_lst):
                        continue

                    tc = tc_lst[col_idx]

                    # 清空被合并单元格的内容
                    for p in tc.p_lst:
                        # 移除所有 run 元素
                        for r in p.r_lst[:]:
                            p.remove(r)

                    # 设置 vMerge 属性为 "continue"
                    merge_tc_pr = tc.get_or_add_tcPr()
                    for old_vmerge in merge_tc_pr.findall(qn('w:vMerge')):
                        merge_tc_pr.remove(old_vmerge)
                    continue_vmerge = OxmlElement('w:vMerge')
                    continue_vmerge.set(qn('w:val'), 'continue')
                    merge_tc_pr.append(continue_vmerge)

                    merged_count += 1
                except Exception as e:
                    print(f"    警告：处理行{row_idx}时出错: {e}")

            print(f"    设置行{group_start+1}-{group_end} vMerge=\"continue\" (共{group_end-group_start}行)")

            # 设置顶部单元格的内容
            # 清空现有内容
            for p in tc_start.p_lst[:]:
                # 移除所有段落（除了保留一个）
                if len(tc_start.p_lst) > 1:
                    tc_start.remove(p)

            # 确保至少有一个段落
            if len(tc_start.p_lst) == 0:
                from docx.oxml import parse_xml
                new_p = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tc_start.append(new_p)

            # 设置内容到第一个段落
            p = tc_start.p_lst[0]
            if merged_text:
                # 清空现有 run
                for r in p.r_lst[:]:
                    p.remove(r)

                # 添加新的 run 和 text
                from docx.oxml import parse_xml
                new_r = parse_xml(r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve"></w:t></w:r>')
                t_element = new_r.find(qn('w:t'))
                if t_element is not None:
                    t_element.text = merged_text
                p.append(new_r)

                content_preview = merged_text[:40] if merged_text else ""
                print(f"    设置行{group_start}内容: \"{content_preview}\"")

                # 验证内容是否设置成功
                verify_text = ""
                for r in p.r_lst:
                    for t in t_lst if (t_lst := r.t_lst) else []:
                        verify_text += t.text
                verify_preview = verify_text[:40] if verify_text else ""
                print(f"    验证行{group_start}内容: \"{verify_preview}\"")
            else:
                print(f"    [警告] merged_text 为空，跳过内容设置")

            # 设置合并后单元格的对齐方式
            try:
                # 验证vMerge属性在设置对齐方式之前是否仍然存在
                tc_pr_check = tc_start.get_or_add_tcPr()
                vmerge_check = tc_pr_check.find(qn('w:vMerge'))
                vmerge_val_check = vmerge_check.get(qn('w:val')) if vmerge_check is not None else 'None'
                print(f"    设置对齐方式前vMerge验证: vMerge=\"{vmerge_val_check}\"")

                # 垂直居中对齐
                tc_pr = tc_start.get_or_add_tcPr()
                v_align = tc_pr.find(qn('w:vAlign'))
                if v_align is None:
                    v_align = OxmlElement('w:vAlign')
                    v_align.set(qn('w:val'), 'center')
                    tc_pr.append(v_align)

                # 段落居中对齐
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                p_pr = p.get_or_add_pPr()
                p_pr.append(jc)

                # 验证vMerge属性在设置对齐方式后是否仍然存在
                vmerge_check2 = tc_pr.find(qn('w:vMerge'))
                vmerge_val_check2 = vmerge_check2.get(qn('w:val')) if vmerge_check2 is not None else 'None'
                print(f"    设置对齐方式后vMerge验证: vMerge=\"{vmerge_val_check2}\"")
            except Exception as e:
                print(f"    设置单元格对齐方式时出错: {e}")

        except Exception as e:
            print(f"    处理合并组时出错: {e}")
            import traceback
            traceback.print_exc()

    print(f"  第 {col_idx} 列纵向合并完成，合并了 {merged_count} 个单元格")
    return merged_count


def _center_table_cells_horizontal(table, start_row=0):
    """将表格所有数据单元格的段落设置水平居中

    Args:
        table: python-docx Table 对象
        start_row: 跳过的表头行数，默认 0（包含表头）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    centered = 0
    for ri, row in enumerate(table.rows):
        if ri < start_row:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p._element.insert(0, pPr)
                if pPr.find(qn('w:jc')) is None:
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'center')
                    pPr.append(jc)
                    centered += 1
    if centered > 0:
        print(f"    已居中对齐 {centered} 个段落")


def merge_other_tables_vertical_cells(doc, context):
    """
    处理范围三类别表格的纵向单元格合并（XML方法）

    仅处理范围三各类别的详细表格，不包含表1和表2。

    Args:
        doc: Word文档对象
        context: 数据上下文字典（用于判断哪些表格有数据）
    """
    total_merged = 0

    # 记录已处理的表格索引，避免重复处理表1和表2
    processed_tables = set()

    # 对有数据的范围三类别表格进行处理
    for cat_num in range(1, 16):
        detail_items = context.get(f'scope3_category{cat_num}', [])
        emission_value = context.get(f'scope_3_category_{cat_num}_emissions', 0)

        # 只处理有数据的类别
        if (detail_items and len(detail_items) > 0) or (emission_value and emission_value > 0):
            category_name = get_scope_3_category_name(cat_num)

            # 查找对应的表格
            table_idx = find_table_by_content(doc, [category_name, f'类别{cat_num}'])

            # 跳过表1和表2（索引0和1）以及已处理的表格
            if table_idx is not None and table_idx < len(doc.tables) and table_idx not in processed_tables:
                # 跳过表1和表2
                if table_idx < 2:
                    print(f"  跳过范围三类别{cat_num}表格（{category_name}，表格索引：{table_idx}），这是主表格")
                    continue

                table = doc.tables[table_idx]
                print(f"  找到范围三类别{cat_num}表格（{category_name}，表格索引：{table_idx}）")
                processed_tables.add(table_idx)

                try:
                    merged = merge_vertical_cells(table, 0)
                    total_merged += merged
                except Exception as e:
                    print(f"  处理范围三类别{cat_num}表格时出错: {e}")
                    import traceback
                    traceback.print_exc()

    print(f"  范围三类别表格合并总计：合并了 {total_merged} 个单元格")
