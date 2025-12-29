#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
对比template.docx和模板1.docx的差异
"""

import docx
import re

def analyze_template_issues(template_file):
    """分析模板的问题"""
    print(f"\n=== 分析 {template_file} ===")

    doc = docx.Document(template_file)

    # 1. 检查语法错误
    syntax_errors = []
    variables = set()

    for paragraph in doc.paragraphs:
        text = paragraph.text
        # 检查所有Jinja2标签
        tags = re.findall(r'(\{\{[^}]+\}\}|\{\%[^%]+\%\})', text)

        for tag in tags:
            variables.add(tag.strip())

            # 检查语法错误
            if '{{' in tag:
                # 检查变量名中是否有空格
                var_content = tag.replace('{{', '').replace('}}', '').strip()
                if ' ' in var_content:
                    syntax_errors.append(f"变量名包含空格: {tag}")
                if var_content.startswith('-') or var_content.endswith('-'):
                    syntax_errors.append(f"变量名格式错误: {tag}")

    # 2. 统计信息
    print(f"段落数量: {len(doc.paragraphs)}")
    print(f"表格数量: {len(doc.tables)}")
    print(f"Jinja2变量数量: {len(variables)}")

    # 3. 显示问题
    if syntax_errors:
        print(f"\n发现 {len(syntax_errors)} 个语法错误:")
        for error in syntax_errors[:5]:  # 只显示前5个
            print(f"  - {error}")
        if len(syntax_errors) > 5:
            print(f"  - ... 还有 {len(syntax_errors) - 5} 个错误")
    else:
        print("\n未发现明显的语法错误")

    # 4. 显示前10个变量
    print(f"\n主要变量（前10个）:")
    for var in sorted(list(variables))[:10]:
        print(f"  - {var}")

    return {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'variables': len(variables),
        'syntax_errors': len(syntax_errors),
        'error_list': syntax_errors
    }

def compare_templates():
    """对比两个模板"""
    print("=== 模板对比分析 ===")

    template1 = analyze_template_issues("template.docx")
    template2 = analyze_template_issues("模板1.docx")

    print("\n=== 对比结果 ===")
    print("项目\t\ttemplate.docx\t模板1.docx")
    print(f"段落数\t\t{template1['paragraphs']}\t\t{template2['paragraphs']}")
    print(f"表格数量\t{template1['tables']}\t\t{template2['tables']}")
    print(f"变量数量\t{template1['variables']}\t\t{template2['variables']}")
    print(f"语法错误\t{template1['syntax_errors']}\t\t{template2['syntax_errors']}")

    print(f"\n=== 结论 ===")
    if template1['syntax_errors'] > 0:
        print("❌ template.docx 存在语法错误，不适合直接使用")
    else:
        print("✅ template.docx 语法正确，可以使用")

    if template2['syntax_errors'] == 0:
        print("✅ 模板1.docx 语法正确，已经验证可以使用")

    if template1['syntax_errors'] > 0 and template2['syntax_errors'] == 0:
        print("\n💡 建议：使用 模板1.docx 作为当前项目的模板")
        print("   如果需要使用 template.docx，需要先修复其中的语法错误")

if __name__ == "__main__":
    compare_templates()