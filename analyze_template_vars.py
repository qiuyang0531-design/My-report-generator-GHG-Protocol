#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析template_fixed.docx中所需的变量
"""

import docx
import re

def analyze_template():
    """分析模板中使用的变量"""
    doc = docx.Document('template_fixed.docx')

    # 提取所有Jinja2变量
    variables = set()

    # 从段落中提取变量
    for paragraph in doc.paragraphs:
        text = paragraph.text
        # 匹配 {{ variable }} 格式的变量
        vars_in_text = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
        for var in vars_in_text:
            variables.add(var.strip())

    # 从表格中提取变量
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    vars_in_text = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                    for var in vars_in_text:
                        variables.add(var.strip())

    # 按字母顺序显示所有变量
    print(f"模板中共需要 {len(variables)} 个变量：")
    print("=" * 50)

    for var in sorted(variables):
        print(f"- {var}")

    # 特别检查包含 'based' 和 'emissions' 的变量
    based_vars = [v for v in variables if 'based' in v]
    emissions_vars = [v for v in variables if 'emissions' in v]

    if based_vars:
        print("\n包含 'based' 的变量：")
        for var in based_vars:
            print(f"- {var}")

    if emissions_vars:
        print("\n包含 'emissions' 的变量：")
        for var in emissions_vars:
            print(f"- {var}")

if __name__ == "__main__":
    analyze_template()