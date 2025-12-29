import docx
from docx.text.paragraph import Paragraph
from docx.table import Table

# 打开模板文件
doc = docx.Document('template.docx')

print("=== 模板检查报告 ===")
print()

# 检查静态内容标签
print("1. 静态内容标签检查：")
static_tags_found = set()
all_static_tags = {'company_name', 'report_year', 'total_emission'}

# 检查段落中的静态标签
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text
    if '{{' in text and '}}' in text:
        print(f"   第{i+1}段：找到标签 - {text.strip()}")
        # 提取所有标签
        import re
        tags = re.findall(r'\{\{\s*([^}\s]+)\s*\}\}', text)
        for tag in tags:
            static_tags_found.add(tag)

# 检查动态表格
print()
print("2. 动态表格检查：")
tables_with_tr_tags = []
tables_with_endtr_tags = []
tables_with_item_tags = []

for i, table in enumerate(doc.tables):
    has_tr_tag = False
    has_endtr_tag = False
    has_item_tags = False
    
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text
            
            if '{% tr for' in cell_text:
                has_tr_tag = True
                tables_with_tr_tags.append(i+1)
                print(f"   表格{i+1}，行{row_idx+1}，单元格{cell_idx+1}：找到表格循环开始标签 - {cell_text.strip()}")
            
            if '{% endtr' in cell_text:
                has_endtr_tag = True
                tables_with_endtr_tags.append(i+1)
                print(f"   表格{i+1}，行{row_idx+1}，单元格{cell_idx+1}：找到表格循环结束标签 - {cell_text.strip()}")
            
            if 'item.' in cell_text and '{{' in cell_text and '}}' in cell_text:
                has_item_tags = True
                tables_with_item_tags.append(i+1)
                print(f"   表格{i+1}，行{row_idx+1}，单元格{cell_idx+1}：找到表格项目标签 - {cell_text.strip()}")
    
    if has_tr_tag or has_endtr_tag or has_item_tags:
        print(f"   表格{i+1}：{'已设置循环标签' if has_tr_tag and has_endtr_tag else '循环标签不完整'}")

# 综合检查结果
print()
print("3. 综合检查结果：")

# 检查静态标签完整性
missing_static_tags = all_static_tags - static_tags_found
if missing_static_tags:
    print(f"   ⚠️  缺少静态标签：{missing_static_tags}")
else:
    print("   ✅ 所有静态标签都已找到")

# 检查动态表格完整性
if tables_with_tr_tags and tables_with_endtr_tags and tables_with_item_tags:
    print("   ✅ 动态表格标签设置基本完整")
    if len(set(tables_with_tr_tags) & set(tables_with_endtr_tags) & set(tables_with_item_tags)) > 0:
        print("   ✅ 存在同时包含开始、结束和项目标签的表格")
    else:
        print("   ⚠️  开始标签、结束标签和项目标签可能不在同一个表格中")
else:
    print("   ⚠️  动态表格标签设置不完整")
    if not tables_with_tr_tags:
        print("     - 未找到表格循环开始标签：{% tr for ... %}")
    if not tables_with_endtr_tags:
        print("     - 未找到表格循环结束标签：{% endtr %}")
    if not tables_with_item_tags:
        print("     - 未找到表格项目标签：{{ item.xxx }}")

print()
print("=== 检查完成 ===")