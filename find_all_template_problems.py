#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
找出template.docx中的所有问题并生成修复方案
"""

import docx
import re

def find_all_problems():
    """找出所有问题"""
    doc = docx.Document('template.docx')

    print("=== template.docx 完整问题清单 ===\n")

    problems = []

    # 扫描所有段落和表格
    all_content = []

    # 段落
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if '{{' in text or '{%' in text:
            all_content.append((f"段落{i+1}", text))

    # 表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if '{{' in text or '{%' in text:
                        location = f"表格{table_idx+1}[{row_idx+1},{col_idx+1}]"
                        all_content.append((location, text))

    print(f"找到 {len(all_content)} 处包含Jinja2标签的内容\n")

    # 分析每个内容
    for location, text in all_content:
        # 提取所有变量
        variables = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
        tags = re.findall(r'(\{\%[^%]+\%\})', text)

        for var in variables:
            original_var = f"{{{{{var}}}}}"
            var_name = var.strip()

            # 检查问题
            if '  ' in var:  # 多个空格
                problems.append((location, original_var, "包含多个空格"))
            elif ' ' in var and not any(op in var for op in ['|', 'filter', 'if']):
                problems.append((location, original_var, "包含空格"))
            elif '-' in var:
                problems.append((location, original_var, "包含连字符"))
            elif var_name != var:  # 首尾有空格
                problems.append((location, original_var, "首尾有空格"))

        for tag in tags:
            if '  ' in tag:
                problems.append((location, tag, "Jinja2标签包含多余空格"))

    # 显示所有问题
    if problems:
        print("发现的问题:")
        current_problem = 1
        for location, tag, issue in problems:
            print(f"{current_problem}. {location}: {tag}")
            print(f"   问题: {issue}")
            current_problem += 1
    else:
        print("未发现问题")

    # 生成修复建议
    print(f"\n=== 修复方案 ===")
    unique_fixes = {}
    for location, tag, issue in problems:
        if '包含连字符' in issue:
            var = re.findall(r'\{\{\s*([^}]+)\s*\}\}', tag)[0].strip()
            if var not in unique_fixes:
                unique_fixes[var] = var.replace('-', '_')
        elif '包含空格' in issue:
            var = re.findall(r'\{\{\s*([^}]+)\s*\}\}', tag)[0]
            if var not in unique_fixes:
                unique_fixes[var] = re.sub(r'\s+', '_', var.strip())

    print(f"需要修复的变量 (共{len(unique_fixes)}个):")
    for original, fixed in sorted(unique_fixes.items()):
        print(f"  {original} → {fixed}")

    # 根据之前错误信息补充可能的问题
    print(f"\n根据运行时错误，还需要注意:")
    print("1. 模板中可能使用了未定义的变量，如 'based_emissions'")
    print("2. 需要确保数据准备代码提供所有模板所需的变量")
    print("3. 某些变量可能在条件语句中被引用")

    return unique_fixes, problems

if __name__ == "__main__":
    fixes, problems = find_all_problems()