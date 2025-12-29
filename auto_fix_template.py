#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
自动修复template.docx中的所有问题
"""

import docx
import re

def fix_template():
    """修复模板中的所有问题"""
    doc = docx.Document('template.docx')

    print("=== 修复template.docx ===\n")

    fix_count = 0

    # 修复规则
    fixes = {
        # 去除尾随空格
        r'(\{\{\s*)(\w+)(\s+\}\})': r'\1\2\3',  # 先捕获
        # 替换连字符为下划线
        r'scope_2_location-based_emissions': 'scope_2_location_based_emissions',
        r'scope_2_market-based_emissions': 'scope_2_market_based_emissions',
    }

    # 处理段落
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        new_text = original_text

        # 去除变量中的尾随空格
        new_text = re.sub(r'\{\{\s*(\w+(?:-\w+)*)\s+\}\}', r'{{\1}}', new_text)

        # 替换连字符
        new_text = new_text.replace('scope_2_location-based_emissions', 'scope_2_location_based_emissions')
        new_text = new_text.replace('scope_2_market-based_emissions', 'scope_2_market_based_emissions')

        if new_text != original_text:
            paragraph.text = new_text
            fix_count += 1
            print(f"修复段落 {i+1}")

    # 处理表格
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text

                    # 去除变量中的尾随空格
                    new_text = re.sub(r'\{\{\s*(\w+(?:-\w+)*)\s+\}\}', r'{{\1}}', new_text)

                    # 替换连字符
                    new_text = new_text.replace('scope_2_location-based_emissions', 'scope_2_location_based_emissions')
                    new_text = new_text.replace('scope_2_market-based_emissions', 'scope_2_market_based_emissions')

                    if new_text != original_text:
                        paragraph.text = new_text
                        fix_count += 1
                        print(f"修复表格 {table_idx+1}[{row_idx+1},{col_idx+1}]")

    # 保存修复后的文件
    doc.save('template_final_fixed.docx')
    print(f"\n修复完成！共修复 {fix_count} 处")
    print("修复后的文件保存为: template_final_fixed.docx")

    return fix_count

if __name__ == "__main__":
    fix_count = fix_template()